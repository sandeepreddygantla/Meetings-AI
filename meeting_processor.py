import os
import json
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, asdict
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import re
from pathlib import Path
# Remove Azure imports and replace with OpenAI
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.schema import HumanMessage, SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
import sqlite3
import faiss
from collections import defaultdict
import threading
import time
import hashlib
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv

# Force reload of environment variables
load_dotenv(override=True)
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# API key will be checked when llm/embedding_model are used, not at module import
logger.info("Meeting processor module loading...")

project_id = "openai-meeting-processor"  # Simple project ID for personal use
tiktoken_cache_dir = os.path.abspath("tiktoken_cache")
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir

# --- Dummy Auth Function (for compatibility) ---
def get_access_token():
    """
    Dummy function to maintain compatibility with existing code.
    OpenAI API uses API key authentication, not tokens.
    """
    return "dummy_token_for_compatibility"

# --- OpenAI LLM Client ---
def get_llm(access_token: str = None):
    """
    Get OpenAI LLM client. access_token parameter is kept for compatibility
    but not used since OpenAI uses API key authentication.
    """
    try:
        # Get fresh API key each time to avoid caching issues
        current_api_key = os.getenv("OPENAI_API_KEY")
        if not current_api_key:
            logger.warning("OPENAI_API_KEY environment variable not set")
            return None
        
        return ChatOpenAI(
            model="gpt-4o",  # Using GPT-4o model
            openai_api_key=current_api_key,
            temperature=0,
            max_tokens=4000,  # Adjust as needed
            request_timeout=60
        )
    except Exception as e:
        logger.error(f"Error creating LLM client: {e}")
        return None

# --- OpenAI Embedding Model ---
def get_embedding_model(access_token: str = None):
    """
    Get OpenAI embedding model. access_token parameter is kept for compatibility
    but not used since OpenAI uses API key authentication.
    """
    try:
        # Get fresh API key each time to avoid caching issues
        current_api_key = os.getenv("OPENAI_API_KEY")
        if not current_api_key:
            logger.warning("OPENAI_API_KEY environment variable not set")
            return None
        
        return OpenAIEmbeddings(
            model="text-embedding-3-large",  # Using text-embedding-3-large
            openai_api_key=current_api_key,
            dimensions=3072  # text-embedding-3-large dimension
        )
    except Exception as e:
        logger.error(f"Error creating embedding model: {e}")
        return None

# Initialize global variables (keeping same structure)
# These will be None if API keys are not available, but won't crash the module load
try:
    access_token = get_access_token()
    embedding_model = get_embedding_model(access_token)
    llm = get_llm(access_token)
    logger.info("LLM and embedding model initialized successfully")
except Exception as e:
    logger.warning(f"Could not initialize LLM/embedding model at startup: {e}")
    access_token = None
    embedding_model = None
    llm = None

@dataclass
class DocumentChunk:
    """Structure to hold document chunk information with enhanced intelligence metadata"""
    chunk_id: str
    document_id: str
    filename: str
    chunk_index: int
    content: str
    start_char: int
    end_char: int
    embedding: Optional[np.ndarray] = None
    # Enhanced intelligence metadata
    enhanced_content: Optional[str] = None
    chunk_type: Optional[str] = None
    speakers: Optional[str] = None  # JSON string
    speaker_contributions: Optional[str] = None  # JSON string
    topics: Optional[str] = None  # JSON string
    decisions: Optional[str] = None  # JSON string
    actions: Optional[str] = None  # JSON string
    questions: Optional[str] = None  # JSON string
    context_before: Optional[str] = None
    context_after: Optional[str] = None
    key_phrases: Optional[str] = None  # JSON string
    importance_score: Optional[float] = None

@dataclass
class User:
    """Structure to hold user information"""
    user_id: str
    username: str
    email: str
    full_name: str
    password_hash: str
    created_at: datetime
    last_login: Optional[datetime] = None
    is_active: bool = True
    role: str = 'user'

@dataclass
class Project:
    """Structure to hold project information"""
    project_id: str
    user_id: str
    project_name: str
    description: str
    created_at: datetime
    is_active: bool = True

@dataclass
class Meeting:
    """Structure to hold meeting information"""
    meeting_id: str
    user_id: str
    project_id: str
    meeting_name: str
    meeting_date: datetime
    created_at: datetime

@dataclass
class MeetingDocument:
    """Structure to hold meeting document information"""
    document_id: str
    filename: str
    date: datetime
    title: str
    content: str
    content_summary: str  # Condensed summary for metadata
    main_topics: List[str]
    past_events: List[str]
    future_actions: List[str]
    participants: List[str]
    chunk_count: int = 0
    file_size: int = 0
    user_id: Optional[str] = None
    meeting_id: Optional[str] = None
    project_id: Optional[str] = None
    folder_path: Optional[str] = None

class VectorDatabase:
    """Vector database using FAISS for similarity search and SQLite for metadata"""
    
    def __init__(self, db_path: str = "meeting_documents.db", index_path: str = "vector_index.faiss"):
        self.db_path = db_path
        self.index_path = index_path
        self.dimension = 3072  # text-embedding-3-large dimension
        self.index = None
        self.chunk_metadata = {}
        self.document_metadata = {}
        self._init_database()
        self._load_or_create_index()
    
    def _init_database(self):
        """Initialize SQLite database for metadata storage"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Create users table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                user_id TEXT PRIMARY KEY,
                username TEXT UNIQUE NOT NULL,
                email TEXT UNIQUE NOT NULL,
                full_name TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                last_login TIMESTAMP,
                is_active BOOLEAN DEFAULT TRUE,
                role TEXT DEFAULT 'user'
            )
        ''')
        
        # Check if we need to migrate existing tables
        self._migrate_existing_tables(cursor)
        
        # Create projects table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS projects (
                project_id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                project_name TEXT NOT NULL,
                description TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                is_active BOOLEAN DEFAULT TRUE,
                FOREIGN KEY (user_id) REFERENCES users (user_id)
            )
        ''')
        
        # Create meetings table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS meetings (
                meeting_id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                project_id TEXT NOT NULL,
                meeting_name TEXT NOT NULL,
                meeting_date DATE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (user_id),
                FOREIGN KEY (project_id) REFERENCES projects (project_id)
            )
        ''')
        
        # Create user sessions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_sessions (
                session_id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                expires_at TIMESTAMP,
                is_active BOOLEAN DEFAULT TRUE,
                FOREIGN KEY (user_id) REFERENCES users (user_id)
            )
        ''')
        
        # Create documents table (updated with user context)
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                document_id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                date TIMESTAMP NOT NULL,
                title TEXT,
                content_summary TEXT,
                main_topics TEXT,
                past_events TEXT,
                future_actions TEXT,
                participants TEXT,
                chunk_count INTEGER,
                file_size INTEGER,
                user_id TEXT,
                meeting_id TEXT,
                project_id TEXT,
                folder_path TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (user_id),
                FOREIGN KEY (meeting_id) REFERENCES meetings (meeting_id),
                FOREIGN KEY (project_id) REFERENCES projects (project_id)
            )
        ''')
        
        # Create chunks table (updated with user context)
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS chunks (
                chunk_id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                content TEXT NOT NULL,
                start_char INTEGER,
                end_char INTEGER,
                user_id TEXT,
                meeting_id TEXT,
                project_id TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (document_id) REFERENCES documents (document_id),
                FOREIGN KEY (user_id) REFERENCES users (user_id),
                FOREIGN KEY (meeting_id) REFERENCES meetings (meeting_id),
                FOREIGN KEY (project_id) REFERENCES projects (project_id)
            )
        ''')
        
        # Create file_hashes table for deduplication
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS file_hashes (
                hash_id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                original_filename TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                file_hash TEXT NOT NULL,
                user_id TEXT NOT NULL,
                project_id TEXT,
                meeting_id TEXT,
                document_id TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (user_id),
                FOREIGN KEY (project_id) REFERENCES projects (project_id),
                FOREIGN KEY (meeting_id) REFERENCES meetings (meeting_id),
                FOREIGN KEY (document_id) REFERENCES documents (document_id),
                UNIQUE(file_hash, user_id)
            )
        ''')
        
        # Create upload_jobs table for tracking background processing
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS upload_jobs (
                job_id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                project_id TEXT,
                meeting_id TEXT,
                total_files INTEGER NOT NULL,
                processed_files INTEGER DEFAULT 0,
                failed_files INTEGER DEFAULT 0,
                status TEXT DEFAULT 'pending',
                error_message TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                started_at TIMESTAMP,
                completed_at TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (user_id),
                FOREIGN KEY (project_id) REFERENCES projects (project_id),
                FOREIGN KEY (meeting_id) REFERENCES meetings (meeting_id)
            )
        ''')
        
        # Create file_processing_status table for individual file tracking
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS file_processing_status (
                status_id TEXT PRIMARY KEY,
                job_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                file_hash TEXT NOT NULL,
                status TEXT DEFAULT 'pending',
                error_message TEXT,
                document_id TEXT,
                chunks_created INTEGER DEFAULT 0,
                processing_started_at TIMESTAMP,
                processing_completed_at TIMESTAMP,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (job_id) REFERENCES upload_jobs (job_id),
                FOREIGN KEY (document_id) REFERENCES documents (document_id)
            )
        ''')
        
        # Create indexes for faster searches
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_document_date ON documents(date)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_document_filename ON documents(filename)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_document_user ON documents(user_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_document_project ON documents(project_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_document_meeting ON documents(meeting_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_chunk_document ON chunks(document_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_chunk_user ON chunks(user_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_users_username ON users(username)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_users_email ON users(email)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_projects_user ON projects(user_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_meetings_user ON meetings(user_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_meetings_project ON meetings(project_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_sessions_user ON user_sessions(user_id)')
        
        conn.commit()
        conn.close()
    
    def _migrate_existing_tables(self, cursor):
        """Migrate existing tables to support multi-user structure and enhanced intelligence"""
        try:
            # Check if documents table exists and needs migration
            cursor.execute("PRAGMA table_info(documents)")
            columns = [column[1] for column in cursor.fetchall()]
            
            if 'user_id' not in columns:
                logger.info("Migrating documents table to support multi-user...")
                cursor.execute('ALTER TABLE documents ADD COLUMN user_id TEXT')
                cursor.execute('ALTER TABLE documents ADD COLUMN meeting_id TEXT')
                cursor.execute('ALTER TABLE documents ADD COLUMN project_id TEXT')
                logger.info("Documents table migrated successfully")
            
            # Check if documents table needs folder_path column
            if 'folder_path' not in columns:
                logger.info("Adding folder_path column to documents table...")
                cursor.execute('ALTER TABLE documents ADD COLUMN folder_path TEXT')
                logger.info("folder_path column added successfully")
            
            # Check if chunks table exists and needs migration
            cursor.execute("PRAGMA table_info(chunks)")
            chunk_columns = [column[1] for column in cursor.fetchall()]
            
            if 'user_id' not in chunk_columns:
                logger.info("Migrating chunks table to support multi-user...")
                cursor.execute('ALTER TABLE chunks ADD COLUMN user_id TEXT')
                cursor.execute('ALTER TABLE chunks ADD COLUMN meeting_id TEXT')
                cursor.execute('ALTER TABLE chunks ADD COLUMN project_id TEXT')
                logger.info("Chunks table migrated successfully")
            
            # Enhanced intelligence metadata migration
            intelligence_columns = [
                'enhanced_content', 'chunk_type', 'speakers', 'speaker_contributions',
                'topics', 'decisions', 'actions', 'questions', 'context_before',
                'context_after', 'key_phrases', 'importance_score'
            ]
            
            missing_intelligence_columns = [col for col in intelligence_columns if col not in chunk_columns]
            
            if missing_intelligence_columns:
                logger.info("Adding enhanced intelligence metadata columns to chunks table...")
                for column in missing_intelligence_columns:
                    if column == 'importance_score':
                        cursor.execute(f'ALTER TABLE chunks ADD COLUMN {column} REAL DEFAULT 0.5')
                    else:
                        cursor.execute(f'ALTER TABLE chunks ADD COLUMN {column} TEXT')
                logger.info(f"Added {len(missing_intelligence_columns)} intelligence metadata columns")
                logger.info("Enhanced intelligence migration completed successfully")
                
        except sqlite3.OperationalError as e:
            # Tables might not exist yet, that's okay
            logger.info(f"Migration check: {e}")
            pass
    
    def _load_or_create_index(self):
        """Load existing FAISS index or create new one"""
        if os.path.exists(self.index_path):
            try:
                self.index = faiss.read_index(self.index_path)
                logger.info(f"Loaded existing FAISS index with {self.index.ntotal} vectors")
                # Rebuild chunk metadata mapping from database
                self._rebuild_chunk_metadata()
            except Exception as e:
                logger.error(f"Error loading FAISS index: {e}")
                self.index = faiss.IndexFlatIP(self.dimension)
        else:
            self.index = faiss.IndexFlatIP(self.dimension)
            logger.info("Created new FAISS index")
    
    def _rebuild_chunk_metadata(self):
        """Rebuild chunk_metadata mapping from database after loading FAISS index"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Get all chunks ordered by their creation (to match FAISS index order)
            cursor.execute('''
                SELECT chunk_id, chunk_index, document_id 
                FROM chunks 
                ORDER BY document_id, chunk_index
            ''')
            
            chunks_data = cursor.fetchall()
            conn.close()
            
            # Rebuild the mapping assuming chunks were added in order
            self.chunk_metadata = {}
            for i, (chunk_id, chunk_index, document_id) in enumerate(chunks_data):
                self.chunk_metadata[i] = chunk_id
            
            logger.info(f"Rebuilt chunk metadata mapping with {len(self.chunk_metadata)} entries")
            
        except Exception as e:
            logger.error(f"Error rebuilding chunk metadata: {e}")
            self.chunk_metadata = {}
    
    def add_document(self, document: MeetingDocument, chunks: List[DocumentChunk]):
        """Add document and its chunks to the database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            # Insert document metadata
            cursor.execute('''
                INSERT OR REPLACE INTO documents 
                (document_id, filename, date, title, content_summary, main_topics, 
                 past_events, future_actions, participants, chunk_count, file_size,
                 user_id, meeting_id, project_id, folder_path)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                document.document_id,
                document.filename,
                document.date,
                document.title,
                document.content_summary,
                json.dumps(document.main_topics),
                json.dumps(document.past_events),
                json.dumps(document.future_actions),
                json.dumps(document.participants),
                document.chunk_count,
                document.file_size,
                document.user_id,
                document.meeting_id,
                document.project_id,
                getattr(document, 'folder_path', None)
            ))
            
            # Prepare vectors and chunk data for batch insertion
            vectors = []
            chunk_data = []
            
            for chunk in chunks:
                if chunk.embedding is not None:
                    vectors.append(chunk.embedding)
                    chunk_data.append((
                        chunk.chunk_id,
                        chunk.document_id,
                        chunk.filename,
                        chunk.chunk_index,
                        chunk.content,
                        chunk.start_char,
                        chunk.end_char
                    ))
            
            # Insert chunks in batch with user context and intelligence metadata
            chunk_data_with_context = []
            for chunk in chunks:
                chunk_data_with_context.append((
                    chunk.chunk_id,
                    chunk.document_id,
                    chunk.filename,
                    chunk.chunk_index,
                    chunk.content,
                    chunk.start_char,
                    chunk.end_char,
                    document.user_id,
                    document.meeting_id,
                    document.project_id,
                    # Enhanced intelligence metadata
                    getattr(chunk, 'enhanced_content', None),
                    getattr(chunk, 'chunk_type', None),
                    getattr(chunk, 'speakers', None),
                    getattr(chunk, 'speaker_contributions', None),
                    getattr(chunk, 'topics', None),
                    getattr(chunk, 'decisions', None),
                    getattr(chunk, 'actions', None),
                    getattr(chunk, 'questions', None),
                    getattr(chunk, 'context_before', None),
                    getattr(chunk, 'context_after', None),
                    getattr(chunk, 'key_phrases', None),
                    getattr(chunk, 'importance_score', 0.5)
                ))
            
            cursor.executemany('''
                INSERT OR REPLACE INTO chunks 
                (chunk_id, document_id, filename, chunk_index, content, start_char, end_char,
                 user_id, meeting_id, project_id, enhanced_content, chunk_type, speakers,
                 speaker_contributions, topics, decisions, actions, questions, context_before,
                 context_after, key_phrases, importance_score)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', chunk_data_with_context)
            
            # Add vectors to FAISS index
            if vectors:
                vectors_array = np.array(vectors).astype('float32')
                # Normalize vectors for cosine similarity
                faiss.normalize_L2(vectors_array)
                self.index.add(vectors_array)
                
                # Store chunk metadata mapping
                start_idx = self.index.ntotal - len(vectors)
                for i, chunk in enumerate(chunks):
                    if chunk.embedding is not None:
                        self.chunk_metadata[start_idx + i] = chunk.chunk_id
            
            conn.commit()
            self.document_metadata[document.document_id] = document
            logger.info(f"Added document {document.filename} with {len(chunks)} chunks")
            
        except Exception as e:
            conn.rollback()
            logger.error(f"Error adding document {document.filename}: {e}")
            raise
        finally:
            conn.close()
    
    def store_document_metadata(self, filename: str, content: str, user_id: str, 
                               project_id: str = None, meeting_id: str = None) -> str:
        """Store document metadata and return document_id"""
        document_id = f"doc_{uuid.uuid4().hex[:8]}"
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            # Insert basic document metadata
            cursor.execute('''
                INSERT INTO documents 
                (document_id, filename, date, title, content_summary, main_topics, 
                 past_events, future_actions, participants, chunk_count, file_size,
                 user_id, meeting_id, project_id, folder_path)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                document_id,
                filename,
                datetime.now(),
                filename,
                "",  # content_summary - will be filled later
                "[]",  # main_topics
                "[]",  # past_events
                "[]",  # future_actions
                "[]",  # participants
                0,  # chunk_count - will be updated later
                len(content.encode('utf-8')),  # file_size
                user_id,
                meeting_id,
                project_id,
                None  # folder_path
            ))
            conn.commit()
            logger.info(f"Stored document metadata for {filename} with ID {document_id}")
            return document_id
            
        except Exception as e:
            conn.rollback()
            logger.error(f"Error storing document metadata for {filename}: {e}")
            raise
        finally:
            conn.close()
    
    def search_similar_chunks(self, query_embedding: np.ndarray, top_k: int = 20) -> List[Tuple[str, float]]:
        """Search for similar chunks using FAISS"""
        if self.index.ntotal == 0:
            return []
        
        # Normalize query vector
        query_vector = query_embedding.reshape(1, -1).astype('float32')
        faiss.normalize_L2(query_vector)
        
        # Search in FAISS index
        similarities, indices = self.index.search(query_vector, min(top_k, self.index.ntotal))
        
        results = []
        for i, idx in enumerate(indices[0]):
            if idx in self.chunk_metadata:
                chunk_id = self.chunk_metadata[idx]
                similarity = float(similarities[0][i])
                results.append((chunk_id, similarity))
        
        return results
    
    def search_similar_chunks_by_folder(self, query_embedding: np.ndarray, user_id: str, folder_path: str, top_k: int = 20) -> List[Tuple[str, float]]:
        """Search for similar chunks using FAISS, filtered by folder"""
        if self.index.ntotal == 0:
            return []
        
        # First get all chunks from semantic search
        all_results = self.search_similar_chunks(query_embedding, top_k * 3)  # Get more to filter
        
        # Filter results by folder
        filtered_results = []
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        for chunk_id, similarity in all_results:
            # Check if this chunk belongs to a document in the specified folder
            cursor.execute('''
                SELECT 1 FROM chunks c
                JOIN documents d ON c.document_id = d.document_id
                WHERE c.chunk_id = ? AND d.user_id = ? AND d.folder_path = ?
            ''', (chunk_id, user_id, folder_path))
            
            if cursor.fetchone():
                filtered_results.append((chunk_id, similarity))
                if len(filtered_results) >= top_k:
                    break
        
        conn.close()
        return filtered_results
    
    def get_chunks_by_ids(self, chunk_ids: List[str]) -> List[DocumentChunk]:
        """Retrieve chunks by their IDs with enhanced intelligence metadata"""
        if not chunk_ids:
            return []
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        placeholders = ','.join(['?' for _ in chunk_ids])
        cursor.execute(f'''
            SELECT chunk_id, document_id, filename, chunk_index, content, start_char, end_char,
                   enhanced_content, chunk_type, speakers, speaker_contributions, topics, 
                   decisions, actions, questions, context_before, context_after, 
                   key_phrases, importance_score
            FROM chunks WHERE chunk_id IN ({placeholders})
        ''', chunk_ids)
        
        chunks = []
        for row in cursor.fetchall():
            chunk = DocumentChunk(
                chunk_id=row[0],
                document_id=row[1],
                filename=row[2],
                chunk_index=row[3],
                content=row[4],
                start_char=row[5],
                end_char=row[6],
                enhanced_content=row[7],
                chunk_type=row[8],
                speakers=row[9],
                speaker_contributions=row[10],
                topics=row[11],
                decisions=row[12],
                actions=row[13],
                questions=row[14],
                context_before=row[15],
                context_after=row[16],
                key_phrases=row[17],
                importance_score=row[18]
            )
            chunks.append(chunk)
        
        conn.close()
        return chunks
    
    def enhanced_search_with_metadata(self, query_embedding: np.ndarray, user_id: str, 
                                    filters: Dict = None, top_k: int = 20) -> List[Dict]:
        """Enhanced search with metadata filtering and context reconstruction"""
        try:
            # Perform base vector search
            base_results = self.search_similar_chunks(query_embedding, top_k * 3)
            
            # Apply metadata filters
            filtered_results = self._apply_metadata_filters(base_results, filters, user_id)
            
            # Get enhanced chunks with metadata
            chunk_ids = [result[0] for result in filtered_results[:top_k]]
            enhanced_chunks = self.get_chunks_by_ids(chunk_ids)
            
            # Reconstruct context around chunks
            context_enhanced_results = []
            for i, chunk in enumerate(enhanced_chunks):
                relevance_score = next((score for cid, score in filtered_results if cid == chunk.chunk_id), 0.0)
                
                context_data = self._reconstruct_chunk_context(chunk)
                
                result = {
                    'chunk': chunk,
                    'relevance_score': relevance_score,
                    'context': context_data,
                    'metadata': {
                        'chunk_type': chunk.chunk_type,
                        'speakers': json.loads(chunk.speakers) if chunk.speakers else [],
                        'topics': json.loads(chunk.topics) if chunk.topics else [],
                        'decisions': json.loads(chunk.decisions) if chunk.decisions else [],
                        'actions': json.loads(chunk.actions) if chunk.actions else [],
                        'importance_score': chunk.importance_score or 0.5
                    }
                }
                context_enhanced_results.append(result)
            
            return context_enhanced_results
            
        except Exception as e:
            logger.error(f"Error in enhanced search: {e}")
            return []
    
    def _apply_metadata_filters(self, search_results: List[Tuple[str, float]], 
                               filters: Dict, user_id: str) -> List[Tuple[str, float]]:
        """Apply metadata filters to search results"""
        if not filters:
            return search_results
        
        filtered_results = []
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            for chunk_id, score in search_results:
                cursor.execute('''
                    SELECT chunk_type, speakers, topics, decisions, actions, importance_score, user_id
                    FROM chunks WHERE chunk_id = ?
                ''', (chunk_id,))
                
                row = cursor.fetchone()
                if not row or row[6] != user_id:  # Check user ownership
                    continue
                
                chunk_type, speakers_json, topics_json, decisions_json, actions_json, importance_score, _ = row
                
                # Apply filters
                if self._chunk_matches_filters(
                    chunk_type, speakers_json, topics_json, decisions_json, 
                    actions_json, importance_score, filters
                ):
                    filtered_results.append((chunk_id, score))
            
        finally:
            conn.close()
        
        return filtered_results
    
    def _chunk_matches_filters(self, chunk_type: str, speakers_json: str, topics_json: str,
                              decisions_json: str, actions_json: str, importance_score: float,
                              filters: Dict) -> bool:
        """Check if chunk matches the specified filters"""
        try:
            # Parse JSON fields
            speakers = json.loads(speakers_json) if speakers_json else []
            topics = json.loads(topics_json) if topics_json else []
            decisions = json.loads(decisions_json) if decisions_json else []
            actions = json.loads(actions_json) if actions_json else []
            
            # Filter by chunk type
            if 'chunk_type' in filters:
                if chunk_type not in filters['chunk_type']:
                    return False
            
            # Filter by speakers (enhanced matching)
            if 'speakers' in filters:
                filter_speakers = [s.lower() for s in filters['speakers']]
                chunk_speakers = [s.lower() for s in speakers]
                
                logger.info(f"Speaker filtering: Looking for {filter_speakers} in chunk speakers {chunk_speakers}")
                
                # Check for exact matches first
                exact_match = any(fs in chunk_speakers for fs in filter_speakers)
                
                # Check for partial matches (first name, last name)
                partial_match = False
                if not exact_match:
                    for fs in filter_speakers:
                        for cs in chunk_speakers:
                            # Check if filter speaker is contained in chunk speaker or vice versa
                            if fs in cs or cs in fs:
                                partial_match = True
                                logger.info(f"Found partial match: '{fs}' matches '{cs}'")
                                break
                            # Check individual name parts
                            fs_parts = fs.split()
                            cs_parts = cs.split()
                            if any(part in cs_parts for part in fs_parts):
                                partial_match = True
                                logger.info(f"Found name part match: '{fs}' parts {fs_parts} match '{cs}' parts {cs_parts}")
                                break
                        if partial_match:
                            break
                
                if not (exact_match or partial_match):
                    logger.info(f"No speaker match found for {filter_speakers} in {chunk_speakers}")
                    return False
                else:
                    logger.info(f"Speaker match found! exact={exact_match}, partial={partial_match}")
            
            # Filter by topics
            if 'topics' in filters:
                filter_topics = [t.lower() for t in filters['topics']]
                chunk_topics = [t.lower() for t in topics]
                if not any(ft in chunk_topics for ft in filter_topics):
                    return False
            
            # Filter by importance score
            if 'min_importance' in filters:
                if (importance_score or 0.5) < filters['min_importance']:
                    return False
            
            # Filter by has_decisions
            if 'has_decisions' in filters and filters['has_decisions']:
                if not decisions:
                    return False
            
            # Filter by has_actions
            if 'has_actions' in filters and filters['has_actions']:
                if not actions:
                    return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error applying filters: {e}")
            return True  # Include chunk if filter check fails
    
    def _reconstruct_chunk_context(self, chunk: DocumentChunk) -> Dict:
        """Reconstruct complete context around a chunk"""
        context = {
            'context_before': chunk.context_before or '',
            'context_after': chunk.context_after or '',
            'enhanced_content': chunk.enhanced_content or chunk.content,
            'related_chunks': [],
            'meeting_context': {},
            'speaker_analysis': {}
        }
        
        try:
            # Get meeting context
            if chunk.document_id:
                meeting_context = self._get_meeting_context(chunk.document_id)
                context['meeting_context'] = meeting_context
            
            # Get related chunks from same document
            related_chunks = self._get_related_chunks(chunk)
            context['related_chunks'] = related_chunks
            
            # Analyze speaker contributions in this chunk
            if chunk.speakers:
                speakers = json.loads(chunk.speakers)
                speaker_contributions = json.loads(chunk.speaker_contributions) if chunk.speaker_contributions else {}
                context['speaker_analysis'] = {
                    'speakers_in_chunk': speakers,
                    'contributions': speaker_contributions
                }
            
        except Exception as e:
            logger.error(f"Error reconstructing context for chunk {chunk.chunk_id}: {e}")
        
        return context
    
    def _get_meeting_context(self, document_id: str) -> Dict:
        """Get meeting-level context for a document"""
        context = {}
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT title, date, main_topics, participants
                FROM documents WHERE document_id = ?
            ''', (document_id,))
            
            row = cursor.fetchone()
            if row:
                context = {
                    'meeting_title': row[0],
                    'meeting_date': row[1],
                    'main_topics': json.loads(row[2]) if row[2] else [],
                    'participants': json.loads(row[3]) if row[3] else []
                }
            
            conn.close()
            
        except Exception as e:
            logger.error(f"Error getting meeting context: {e}")
        
        return context
    
    def _get_related_chunks(self, chunk: DocumentChunk, max_related: int = 3) -> List[Dict]:
        """Get related chunks from the same document"""
        related = []
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Get chunks from same document near this chunk
            cursor.execute('''
                SELECT chunk_id, chunk_index, content, chunk_type, importance_score
                FROM chunks 
                WHERE document_id = ? AND chunk_id != ?
                ORDER BY ABS(chunk_index - ?) ASC
                LIMIT ?
            ''', (chunk.document_id, chunk.chunk_id, chunk.chunk_index, max_related))
            
            for row in cursor.fetchall():
                related.append({
                    'chunk_id': row[0],
                    'chunk_index': row[1],
                    'content_preview': row[2][:200] + "..." if len(row[2]) > 200 else row[2],
                    'chunk_type': row[3],
                    'importance_score': row[4]
                })
            
            conn.close()
            
        except Exception as e:
            logger.error(f"Error getting related chunks: {e}")
        
        return related
    
    def get_documents_by_timeframe(self, timeframe: str, user_id: str = None) -> List[MeetingDocument]:
        """Get documents filtered by intelligent timeframe calculation"""
        
        # Calculate date range using intelligent calendar logic
        start_date, end_date = self._calculate_date_range(timeframe)
        
        if not start_date and not end_date:
            logger.warning(f"Unknown timeframe: {timeframe}, returning all documents")
            return list(self.document_metadata.values())
        
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Build query with user filtering if provided
        base_query = '''
            SELECT document_id, filename, date, title, content_summary, main_topics,
                   past_events, future_actions, participants, chunk_count, file_size,
                   user_id, project_id, folder_path
            FROM documents 
            WHERE 1=1
        '''
        
        params = []
        
        # Add user filtering
        if user_id:
            base_query += ' AND user_id = ?'
            params.append(user_id)
        
        # Add date filtering
        if start_date and end_date:
            base_query += ' AND date BETWEEN ? AND ?'
            params.extend([start_date.isoformat(), end_date.isoformat()])
        elif start_date:
            base_query += ' AND date >= ?'
            params.append(start_date.isoformat())
        elif end_date:
            base_query += ' AND date <= ?'
            params.append(end_date.isoformat())
        
        base_query += ' ORDER BY date DESC'
        
        cursor.execute(base_query, params)
        
        documents = []
        for row in cursor.fetchall():
            doc = MeetingDocument(
                document_id=row[0],
                filename=row[1],
                date=datetime.fromisoformat(row[2]),
                title=row[3],
                content="",  # We don't load full content for filtering
                content_summary=row[4],
                main_topics=json.loads(row[5]) if row[5] else [],
                past_events=json.loads(row[6]) if row[6] else [],
                future_actions=json.loads(row[7]) if row[7] else [],
                participants=json.loads(row[8]) if row[8] else [],
                chunk_count=row[9],
                file_size=row[10]
            )
            documents.append(doc)
        
        conn.close()
        return documents
    
    def _calculate_date_range(self, timeframe: str) -> Tuple[Optional[datetime], Optional[datetime]]:
        """Calculate start and end dates for intelligent timeframe filtering"""
        import calendar
        
        now = datetime.now()
        today = now.date()
        
        # Get current week boundaries (Monday to Sunday)
        current_week_start = today - timedelta(days=today.weekday())
        current_week_end = current_week_start + timedelta(days=6)
        
        # Get last week boundaries
        last_week_start = current_week_start - timedelta(days=7)
        last_week_end = current_week_start - timedelta(days=1)
        
        # Get current month boundaries
        current_month_start = today.replace(day=1)
        _, last_day = calendar.monthrange(today.year, today.month)
        current_month_end = today.replace(day=last_day)
        
        # Get last month boundaries
        if today.month == 1:
            last_month_year = today.year - 1
            last_month = 12
        else:
            last_month_year = today.year
            last_month = today.month - 1
        
        last_month_start = today.replace(year=last_month_year, month=last_month, day=1)
        _, last_month_last_day = calendar.monthrange(last_month_year, last_month)
        last_month_end = today.replace(year=last_month_year, month=last_month, day=last_month_last_day)
        
        # Get quarter boundaries
        current_quarter = ((today.month - 1) // 3) + 1
        current_quarter_start = today.replace(month=(current_quarter - 1) * 3 + 1, day=1)
        quarter_end_month = current_quarter * 3
        _, quarter_last_day = calendar.monthrange(today.year, quarter_end_month)
        current_quarter_end = today.replace(month=quarter_end_month, day=quarter_last_day)
        
        # Map timeframes to date ranges
        timeframe_map = {
            # Current periods
            'current_week': (datetime.combine(current_week_start, datetime.min.time()), 
                           datetime.combine(current_week_end, datetime.max.time())),
            'this_week': (datetime.combine(current_week_start, datetime.min.time()), 
                         datetime.combine(current_week_end, datetime.max.time())),
            'current_month': (datetime.combine(current_month_start, datetime.min.time()), 
                            datetime.combine(current_month_end, datetime.max.time())),
            'this_month': (datetime.combine(current_month_start, datetime.min.time()), 
                         datetime.combine(current_month_end, datetime.max.time())),
            
            # Last periods
            'last_week': (datetime.combine(last_week_start, datetime.min.time()), 
                         datetime.combine(last_week_end, datetime.max.time())),
            'past_week': (datetime.combine(last_week_start, datetime.min.time()), 
                         datetime.combine(last_week_end, datetime.max.time())),
            'last_month': (datetime.combine(last_month_start, datetime.min.time()), 
                          datetime.combine(last_month_end, datetime.max.time())),
            'past_month': (datetime.combine(last_month_start, datetime.min.time()), 
                          datetime.combine(last_month_end, datetime.max.time())),
            
            # Relative periods (from start date to now)
            'last_7_days': (now - timedelta(days=7), now),
            'past_7_days': (now - timedelta(days=7), now),
            'last_14_days': (now - timedelta(days=14), now),
            'last_30_days': (now - timedelta(days=30), now),
            'last_60_days': (now - timedelta(days=60), now),
            'last_90_days': (now - timedelta(days=90), now),
            
            # Recent periods
            'recent': (now - timedelta(days=30), now),
            'recently': (now - timedelta(days=30), now),
            
            # Quarter periods
            'current_quarter': (datetime.combine(current_quarter_start, datetime.min.time()), 
                              datetime.combine(current_quarter_end, datetime.max.time())),
            'this_quarter': (datetime.combine(current_quarter_start, datetime.min.time()), 
                           datetime.combine(current_quarter_end, datetime.max.time())),
            'last_quarter': (None, None),  # Will be calculated below
            
            # Year periods
            'current_year': (datetime(today.year, 1, 1), datetime(today.year, 12, 31, 23, 59, 59)),
            'this_year': (datetime(today.year, 1, 1), datetime(today.year, 12, 31, 23, 59, 59)),
            'last_year': (datetime(today.year - 1, 1, 1), datetime(today.year - 1, 12, 31, 23, 59, 59)),
            
            # Extended periods
            'last_3_months': (now - timedelta(days=90), now),
            'last_6_months': (now - timedelta(days=180), now),
            'last_12_months': (now - timedelta(days=365), now),
        }
        
        # Handle last quarter calculation
        if timeframe.lower() in ['last_quarter', 'past_quarter']:
            if current_quarter == 1:
                last_quarter = 4
                last_quarter_year = today.year - 1
            else:
                last_quarter = current_quarter - 1
                last_quarter_year = today.year
            
            last_quarter_start = datetime(last_quarter_year, (last_quarter - 1) * 3 + 1, 1)
            quarter_end_month = last_quarter * 3
            _, quarter_last_day = calendar.monthrange(last_quarter_year, quarter_end_month)
            last_quarter_end = datetime(last_quarter_year, quarter_end_month, quarter_last_day, 23, 59, 59)
            timeframe_map['last_quarter'] = (last_quarter_start, last_quarter_end)
            timeframe_map['past_quarter'] = (last_quarter_start, last_quarter_end)
        
        # Get the timeframe, case-insensitive
        timeframe_key = timeframe.lower().replace(' ', '_')
        if timeframe_key in timeframe_map:
            return timeframe_map[timeframe_key]
        
        # If no exact match, try partial matches
        for key in timeframe_map:
            if timeframe_key in key or key in timeframe_key:
                logger.info(f"Partial match found: {timeframe_key} -> {key}")
                return timeframe_map[key]
        
        logger.warning(f"No date range found for timeframe: {timeframe}")
        return None, None
    
    def _generate_date_based_summary(self, query: str, documents: List[Any], timeframe: str, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Generate intelligent date-based summary with chronological organization"""
        
        # Sort documents by date
        sorted_docs = sorted(documents, key=lambda x: x.date)
        
        # Group documents by date for better organization
        from collections import defaultdict
        date_groups = defaultdict(list)
        for doc in sorted_docs:
            date_key = doc.date.strftime('%Y-%m-%d')
            date_groups[date_key].append(doc)
        
        # Build comprehensive context from all documents
        context_parts = []
        document_summaries = []
        
        for date_key, docs in sorted(date_groups.items()):
            date_formatted = datetime.strptime(date_key, '%Y-%m-%d').strftime('%B %d, %Y')
            context_parts.append(f"\n=== {date_formatted} ===")
            
            for doc in docs:
                # Add document summary
                doc_summary = f"Document: {doc.filename}\n"
                if doc.content_summary:
                    doc_summary += f"Summary: {doc.content_summary}\n"
                if doc.main_topics:
                    doc_summary += f"Main Topics: {', '.join(doc.main_topics)}\n"
                if doc.participants:
                    doc_summary += f"Participants: {', '.join(doc.participants)}\n"
                if doc.future_actions:
                    doc_summary += f"Action Items: {', '.join(doc.future_actions)}\n"
                
                context_parts.append(doc_summary)
                document_summaries.append(doc_summary)
        
        # Create comprehensive context
        full_context = '\n'.join(context_parts)
        
        # Generate summary prompt based on query type
        timeframe_display = timeframe.replace('_', ' ').title()
        summary_prompt = f"""
        The user asked: "{query}"
        
        Based on the meeting documents from {timeframe_display}, please answer their question naturally and comprehensively. Focus on what they specifically asked for rather than forcing a predetermined structure.

        Meeting Documents Context:
        {full_context}

        Provide a direct, helpful answer that addresses exactly what the user wants to know about {timeframe_display}.
        
        IMPORTANT: When referencing information, always cite the specific document filename rather than document numbers or chunk references.
        """
        
        try:
            # Use class LLM instance
            messages = [
                SystemMessage(content="You are a helpful AI assistant that answers questions about meeting documents naturally. Focus on what the user specifically asked for rather than forcing a predetermined structure. Always cite document filenames rather than chunk numbers when referencing information."),
                HumanMessage(content=summary_prompt)
            ]
            
            response = self.llm.invoke(messages)
            summary_response = response.content.strip()
            
            # Add timeframe context to the response
            final_response = f"**Summary for {timeframe_display}** ({len(documents)} documents)\n\n{summary_response}"
            
            if include_context:
                return final_response, full_context
            else:
                return final_response
                
        except Exception as e:
            logger.error(f"Error generating date-based summary: {e}")
            # Fallback to basic summary
            fallback_summary = f"""**Summary for {timeframe_display}** ({len(documents)} documents)

            During this period, there were {len(documents)} meetings/documents:
            
            {chr(10).join([f"• {doc.filename} ({doc.date.strftime('%B %d, %Y')})" for doc in sorted_docs])}
            
            For detailed information about specific meetings, please ask about individual documents or topics.
            """
            
            if include_context:
                return fallback_summary, full_context
            else:
                return fallback_summary
    
    def keyword_search_chunks(self, keywords: List[str], limit: int = 50) -> List[str]:
        """Perform keyword search on chunk content"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Build search query
        search_conditions = []
        params = []
        for keyword in keywords:
            search_conditions.append("content LIKE ?")
            params.append(f"%{keyword}%")
        
        query = f'''
            SELECT chunk_id FROM chunks 
            WHERE {" OR ".join(search_conditions)}
            LIMIT ?
        '''
        params.append(limit)
        
        cursor.execute(query, params)
        chunk_ids = [row[0] for row in cursor.fetchall()]
        
        conn.close()
        return chunk_ids
    
    def get_all_documents(self, user_id: str = None) -> List[Dict[str, Any]]:
        """Get all documents with metadata for document selection"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        if user_id:
            cursor.execute('''
                SELECT d.document_id, d.filename, d.date, d.title, d.content_summary, d.file_size, d.chunk_count,
                       d.user_id, d.meeting_id, d.project_id, d.folder_path, p.project_name
                FROM documents d
                LEFT JOIN projects p ON d.project_id = p.project_id
                WHERE d.user_id = ?
                ORDER BY d.date DESC
            ''', (user_id,))
        else:
            cursor.execute('''
                SELECT d.document_id, d.filename, d.date, d.title, d.content_summary, d.file_size, d.chunk_count,
                       d.user_id, d.meeting_id, d.project_id, d.folder_path, p.project_name
                FROM documents d
                LEFT JOIN projects p ON d.project_id = p.project_id
                ORDER BY d.date DESC
            ''')
        
        documents = []
        for row in cursor.fetchall():
            doc = {
                'document_id': row[0],
                'filename': row[1],
                'date': row[2],
                'title': row[3],
                'content_summary': row[4],
                'file_size': row[5],
                'chunk_count': row[6],
                'user_id': row[7],
                'meeting_id': row[8],
                'project_id': row[9],
                'folder_path': row[10],
                'project_name': row[11]
            }
            documents.append(doc)
        
        conn.close()
        return documents
    
    def get_user_documents_by_scope(self, user_id: str, project_id: str = None, meeting_id: Union[str, List[str]] = None) -> List[str]:
        """Get document IDs for a user filtered by project or meeting(s)"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        if meeting_id:
            if isinstance(meeting_id, list):
                # Handle multiple meeting IDs
                placeholders = ','.join(['?' for _ in meeting_id])
                cursor.execute(f'''
                    SELECT document_id FROM documents 
                    WHERE user_id = ? AND meeting_id IN ({placeholders})
                ''', [user_id] + meeting_id)
            else:
                # Handle single meeting ID
                cursor.execute('''
                    SELECT document_id FROM documents 
                    WHERE user_id = ? AND meeting_id = ?
                ''', (user_id, meeting_id))
        elif project_id:
            cursor.execute('''
                SELECT document_id FROM documents 
                WHERE user_id = ? AND project_id = ?
            ''', (user_id, project_id))
        else:
            cursor.execute('''
                SELECT document_id FROM documents 
                WHERE user_id = ?
            ''', (user_id,))
        
        document_ids = [row[0] for row in cursor.fetchall()]
        conn.close()
        return document_ids
    
    def get_user_documents_by_folder(self, user_id: str, folder_path: str, project_id: str = None, meeting_id: Union[str, List[str]] = None) -> List[str]:
        """Get document IDs for a user filtered by folder path"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # When folder_path is specified, ONLY filter by user and folder
        # This ensures we get ALL documents in the specified folder
        conditions = ["user_id = ?", "folder_path = ?"]
        params = [user_id, folder_path]
        
        # Only add additional filters if folder_path is not the primary filter
        # This prevents conflicting conditions when user selects a specific folder
        if project_id and meeting_id:
            # Both project and meeting specified - this is likely an error case
            # but we'll handle it by prioritizing folder_path only
            logger.warning(f"Both project_id and meeting_id specified with folder_path - using folder_path only")
        
        cursor.execute(f'''
            SELECT document_id FROM documents 
            WHERE {' AND '.join(conditions)}
        ''', params)
        
        document_ids = [row[0] for row in cursor.fetchall()]
        conn.close()
        
        logger.info(f"Found {len(document_ids)} documents in folder {folder_path} for user {user_id}")
        return document_ids
    
    def keyword_search_chunks_by_user(self, keywords: List[str], user_id: str, project_id: str = None, meeting_id: Union[str, List[str]] = None, limit: int = 50) -> List[str]:
        """Perform keyword search on chunk content filtered by user/project/meeting"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Build search conditions
        search_conditions = []
        params = []
        
        # Add keyword conditions
        for keyword in keywords:
            search_conditions.append("content LIKE ?")
            params.append(f"%{keyword}%")
        
        # Add user filter
        user_condition = "user_id = ?"
        params.append(user_id)
        
        # Add optional project/meeting filters
        additional_conditions = [user_condition]
        if project_id:
            additional_conditions.append("project_id = ?")
            params.append(project_id)
        if meeting_id:
            if isinstance(meeting_id, list):
                # Handle multiple meeting IDs
                placeholders = ','.join(['?' for _ in meeting_id])
                additional_conditions.append(f"meeting_id IN ({placeholders})")
                params.extend(meeting_id)
            else:
                # Handle single meeting ID
                additional_conditions.append("meeting_id = ?")
                params.append(meeting_id)
        
        where_clause = f"({' OR '.join(search_conditions)}) AND {' AND '.join(additional_conditions)}"
        
        query = f'''
            SELECT chunk_id FROM chunks 
            WHERE {where_clause}
            LIMIT ?
        '''
        params.append(limit)
        
        cursor.execute(query, params)
        chunk_ids = [row[0] for row in cursor.fetchall()]
        
        conn.close()
        return chunk_ids
    
    def keyword_search_chunks_by_folder(self, keywords: List[str], user_id: str, folder_path: str, limit: int = 50) -> List[str]:
        """Perform keyword search on chunk content filtered by user and folder"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Build search conditions
        search_conditions = []
        params = []
        
        # Add keyword conditions
        for keyword in keywords:
            search_conditions.append("c.content LIKE ?")
            params.append(f"%{keyword}%")
        
        # Add user and folder filters
        where_clause = f"({' OR '.join(search_conditions)}) AND d.user_id = ? AND d.folder_path = ?"
        params.extend([user_id, folder_path])
        
        query = f'''
            SELECT c.chunk_id FROM chunks c
            JOIN documents d ON c.document_id = d.document_id
            WHERE {where_clause}
            LIMIT ?
        '''
        params.append(limit)
        
        cursor.execute(query, params)
        chunk_ids = [row[0] for row in cursor.fetchall()]
        
        conn.close()
        return chunk_ids
    
    # User Management Methods
    def create_user(self, username: str, email: str, full_name: str, password_hash: str) -> str:
        """Create a new user"""
        user_id = f"user_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{username}"
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT INTO users (user_id, username, email, full_name, password_hash)
                VALUES (?, ?, ?, ?, ?)
            ''', (user_id, username, email, full_name, password_hash))
            
            conn.commit()
            logger.info(f"Created user: {username} ({user_id})")
            return user_id
            
        except sqlite3.IntegrityError as e:
            logger.error(f"Error creating user {username}: {e}")
            conn.rollback()
            raise ValueError(f"Username or email already exists")
        finally:
            conn.close()
    
    def get_user_by_username(self, username: str) -> Optional[User]:
        """Get user by username"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT user_id, username, email, full_name, password_hash, created_at, last_login, is_active, role
            FROM users WHERE username = ? AND is_active = TRUE
        ''', (username,))
        
        row = cursor.fetchone()
        conn.close()
        
        if row:
            return User(
                user_id=row[0],
                username=row[1],
                email=row[2],
                full_name=row[3],
                password_hash=row[4],
                created_at=datetime.fromisoformat(row[5]),
                last_login=datetime.fromisoformat(row[6]) if row[6] else None,
                is_active=bool(row[7]),
                role=row[8]
            )
        return None
    
    def get_user_by_id(self, user_id: str) -> Optional[User]:
        """Get user by user_id"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT user_id, username, email, full_name, password_hash, created_at, last_login, is_active, role
            FROM users WHERE user_id = ? AND is_active = TRUE
        ''', (user_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if row:
            return User(
                user_id=row[0],
                username=row[1],
                email=row[2],
                full_name=row[3],
                password_hash=row[4],
                created_at=datetime.fromisoformat(row[5]),
                last_login=datetime.fromisoformat(row[6]) if row[6] else None,
                is_active=bool(row[7]),
                role=row[8]
            )
        return None
    
    def update_user_last_login(self, user_id: str):
        """Update user's last login timestamp"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE user_id = ?
        ''', (user_id,))
        
        conn.commit()
        conn.close()
    
    # Project Management Methods
    def create_project(self, user_id: str, project_name: str, description: str = "") -> str:
        """Create a new project for a user"""
        project_id = f"proj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{user_id.split('_')[-1]}"
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT INTO projects (project_id, user_id, project_name, description)
                VALUES (?, ?, ?, ?)
            ''', (project_id, user_id, project_name, description))
            
            conn.commit()
            logger.info(f"Created project: {project_name} for user {user_id}")
            return project_id
            
        except Exception as e:
            logger.error(f"Error creating project {project_name}: {e}")
            conn.rollback()
            raise
        finally:
            conn.close()
    
    def get_user_projects(self, user_id: str) -> List[Project]:
        """Get all projects for a user"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT project_id, user_id, project_name, description, created_at, is_active
            FROM projects WHERE user_id = ? AND is_active = TRUE
            ORDER BY created_at DESC
        ''', (user_id,))
        
        projects = []
        for row in cursor.fetchall():
            project = Project(
                project_id=row[0],
                user_id=row[1],
                project_name=row[2],
                description=row[3],
                created_at=datetime.fromisoformat(row[4]),
                is_active=bool(row[5])
            )
            projects.append(project)
        
        conn.close()
        return projects
    
    # Meeting Management Methods
    def create_meeting(self, user_id: str, project_id: str, meeting_name: str, meeting_date: datetime) -> str:
        """Create a new meeting"""
        meeting_id = f"meet_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{user_id.split('_')[-1]}"
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT INTO meetings (meeting_id, user_id, project_id, meeting_name, meeting_date)
                VALUES (?, ?, ?, ?, ?)
            ''', (meeting_id, user_id, project_id, meeting_name, meeting_date.date()))
            
            conn.commit()
            logger.info(f"Created meeting: {meeting_name} for project {project_id}")
            return meeting_id
            
        except Exception as e:
            logger.error(f"Error creating meeting {meeting_name}: {e}")
            conn.rollback()
            raise
        finally:
            conn.close()
    
    def get_user_meetings(self, user_id: str, project_id: str = None) -> List[Meeting]:
        """Get meetings for a user, optionally filtered by project"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        if project_id:
            cursor.execute('''
                SELECT meeting_id, user_id, project_id, meeting_name, meeting_date, created_at
                FROM meetings WHERE user_id = ? AND project_id = ?
                ORDER BY meeting_date DESC, created_at DESC
            ''', (user_id, project_id))
        else:
            cursor.execute('''
                SELECT meeting_id, user_id, project_id, meeting_name, meeting_date, created_at
                FROM meetings WHERE user_id = ?
                ORDER BY meeting_date DESC, created_at DESC
            ''', (user_id,))
        
        meetings = []
        for row in cursor.fetchall():
            meeting = Meeting(
                meeting_id=row[0],
                user_id=row[1],
                project_id=row[2],
                meeting_name=row[3],
                meeting_date=datetime.fromisoformat(row[4]),
                created_at=datetime.fromisoformat(row[5])
            )
            meetings.append(meeting)
        
        conn.close()
        return meetings
    
    def get_all_documents(self, user_id: str) -> List[Dict[str, Any]]:
        """Get all documents for a user"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT document_id, filename, date, title, content_summary, 
                   main_topics, past_events, future_actions, participants,
                   chunk_count, file_size, user_id, meeting_id, project_id, folder_path
            FROM documents 
            WHERE user_id = ?
            ORDER BY date DESC
        ''', (user_id,))
        
        documents = []
        for row in cursor.fetchall():
            doc_dict = {
                'document_id': row[0],
                'filename': row[1],
                'date': row[2],
                'title': row[3],
                'content_summary': row[4],
                'main_topics': json.loads(row[5] or '[]'),
                'past_events': json.loads(row[6] or '[]'),
                'future_actions': json.loads(row[7] or '[]'),
                'participants': json.loads(row[8] or '[]'),
                'chunk_count': row[9],
                'file_size': row[10],
                'user_id': row[11],
                'meeting_id': row[12],
                'project_id': row[13],
                'folder_path': row[14]
            }
            documents.append(doc_dict)
        
        conn.close()
        return documents
    
    def get_project_documents(self, project_id: str, user_id: str) -> List[Dict[str, Any]]:
        """Get all documents for a specific project"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT document_id, filename, date, title, content_summary, 
                   main_topics, past_events, future_actions, participants,
                   chunk_count, file_size, user_id, meeting_id, project_id, folder_path
            FROM documents 
            WHERE project_id = ? AND user_id = ?
            ORDER BY date DESC
        ''', (project_id, user_id))
        
        documents = []
        for row in cursor.fetchall():
            doc_dict = {
                'document_id': row[0],
                'filename': row[1],
                'date': row[2],
                'title': row[3],
                'content_summary': row[4],
                'main_topics': json.loads(row[5] or '[]'),
                'past_events': json.loads(row[6] or '[]'),
                'future_actions': json.loads(row[7] or '[]'),
                'participants': json.loads(row[8] or '[]'),
                'chunk_count': row[9],
                'file_size': row[10],
                'user_id': row[11],
                'meeting_id': row[12],
                'project_id': row[13],
                'folder_path': row[14]
            }
            documents.append(doc_dict)
        
        conn.close()
        return documents
    
    def get_document_metadata(self, document_id: str) -> Dict[str, Any]:
        """Get metadata for a specific document"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT document_id, filename, date, title, content_summary, 
                   main_topics, past_events, future_actions, participants,
                   chunk_count, file_size, user_id, meeting_id, project_id
            FROM documents 
            WHERE document_id = ?
        ''', (document_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if row:
            return {
                'document_id': row[0],
                'filename': row[1],
                'date': row[2],
                'title': row[3],
                'content_summary': row[4],
                'main_topics': json.loads(row[5] or '[]'),
                'past_events': json.loads(row[6] or '[]'),
                'future_actions': json.loads(row[7] or '[]'),
                'participants': json.loads(row[8] or '[]'),
                'chunk_count': row[9],
                'file_size': row[10],
                'user_id': row[11],
                'meeting_id': row[12],
                'project_id': row[13]
            }
        return None
    
    def create_session(self, user_id: str, session_id: str, expires_at: datetime) -> bool:
        """Create a new user session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT INTO user_sessions (session_id, user_id, expires_at, is_active)
                VALUES (?, ?, ?, ?)
            ''', (session_id, user_id, expires_at.isoformat(), True))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            logger.error(f"Error creating session: {e}")
            return False
    
    def validate_session(self, session_id: str) -> Optional[str]:
        """Validate a session and return user_id if valid"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT user_id, expires_at FROM user_sessions 
                WHERE session_id = ? AND is_active = TRUE
            ''', (session_id,))
            
            result = cursor.fetchone()
            conn.close()
            
            if result:
                user_id, expires_at = result
                # Check if session is still valid
                expiry_time = datetime.fromisoformat(expires_at)
                if expiry_time > datetime.now():
                    return user_id
                else:
                    # Session expired, deactivate it
                    self.deactivate_session(session_id)
                    return None
            
            return None
            
        except Exception as e:
            logger.error(f"Error validating session: {e}")
            return None
    
    def deactivate_session(self, session_id: str) -> bool:
        """Deactivate a session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                UPDATE user_sessions SET is_active = FALSE 
                WHERE session_id = ?
            ''', (session_id,))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            logger.error(f"Error deactivating session: {e}")
            return False
    
    def cleanup_expired_sessions(self) -> int:
        """Clean up expired sessions and return count of cleaned sessions"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            now = datetime.now().isoformat()
            
            # Count expired sessions
            cursor.execute('''
                SELECT COUNT(*) FROM user_sessions 
                WHERE expires_at < ? AND is_active = TRUE
            ''', (now,))
            
            count = cursor.fetchone()[0]
            
            # Deactivate expired sessions
            cursor.execute('''
                UPDATE user_sessions SET is_active = FALSE 
                WHERE expires_at < ? AND is_active = TRUE
            ''', (now,))
            
            conn.commit()
            conn.close()
            
            if count > 0:
                logger.info(f"Cleaned up {count} expired sessions")
            
            return count
            
        except Exception as e:
            logger.error(f"Error cleaning up expired sessions: {e}")
            return 0
    
    def extend_session(self, session_id: str, new_expires_at: datetime) -> bool:
        """Extend a session's expiry time"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                UPDATE user_sessions SET expires_at = ? 
                WHERE session_id = ? AND is_active = TRUE
            ''', (new_expires_at.isoformat(), session_id))
            
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            logger.error(f"Error extending session: {e}")
            return False
    
    def save_index(self):
        """Save FAISS index to disk"""
        if self.index:
            faiss.write_index(self.index, self.index_path)
            logger.info(f"Saved FAISS index with {self.index.ntotal} vectors")
    
    # File Deduplication Methods
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of a file"""
        hasher = hashlib.sha256()
        try:
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating hash for {file_path}: {e}")
            raise

    def is_file_duplicate(self, file_hash: str, filename: str, user_id: str) -> Optional[Dict]:
        """Check if file is a duplicate based on hash and return original file info"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                SELECT hash_id, filename, original_filename, created_at, document_id
                FROM file_hashes 
                WHERE file_hash = ? AND user_id = ?
                ORDER BY created_at DESC
                LIMIT 1
            ''', (file_hash, user_id))
            
            result = cursor.fetchone()
            if result:
                return {
                    'hash_id': result[0],
                    'filename': result[1],
                    'original_filename': result[2],
                    'created_at': result[3],
                    'document_id': result[4]
                }
            return None
        finally:
            conn.close()

    def store_file_hash(self, file_hash: str, filename: str, original_filename: str, 
                       file_size: int, user_id: str, project_id: str = None, 
                       meeting_id: str = None, document_id: str = None) -> str:
        """Store file hash for deduplication"""
        hash_id = f"hash_{uuid.uuid4().hex[:8]}"
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT INTO file_hashes 
                (hash_id, filename, original_filename, file_size, file_hash, 
                 user_id, project_id, meeting_id, document_id)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (hash_id, filename, original_filename, file_size, file_hash, 
                  user_id, project_id, meeting_id, document_id))
            conn.commit()
            return hash_id
        except sqlite3.IntegrityError:
            logger.warning(f"File hash {file_hash} already exists for user {user_id}")
            raise
        finally:
            conn.close()

    # Background Processing Methods
    def create_upload_job(self, user_id: str, total_files: int, project_id: str = None, 
                         meeting_id: str = None) -> str:
        """Create a new upload job for tracking progress"""
        job_id = f"job_{uuid.uuid4().hex[:8]}"
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT INTO upload_jobs 
                (job_id, user_id, project_id, meeting_id, total_files, status)
                VALUES (?, ?, ?, ?, ?, 'pending')
            ''', (job_id, user_id, project_id, meeting_id, total_files))
            conn.commit()
            logger.info(f"Created upload job {job_id} for {total_files} files")
            return job_id
        finally:
            conn.close()

    def get_job_status(self, job_id: str) -> Optional[Dict]:
        """Get current job status and progress"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                SELECT job_id, user_id, project_id, meeting_id, total_files, 
                       processed_files, failed_files, status, error_message,
                       created_at, started_at, completed_at
                FROM upload_jobs 
                WHERE job_id = ?
            ''', (job_id,))
            
            result = cursor.fetchone()
            if result:
                return {
                    'job_id': result[0],
                    'user_id': result[1],
                    'project_id': result[2],
                    'meeting_id': result[3],
                    'total_files': result[4],
                    'processed_files': result[5],
                    'failed_files': result[6],
                    'status': result[7],
                    'error_message': result[8],
                    'created_at': result[9],
                    'started_at': result[10],
                    'completed_at': result[11],
                    'progress_percent': (result[5] / result[4]) * 100 if result[4] > 0 else 0
                }
            return None
        finally:
            conn.close()

    def update_job_status(self, job_id: str, status: str, processed_files: int = None, 
                         failed_files: int = None, error_message: str = None):
        """Update job status and progress"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            update_fields = ['status = ?']
            params = [status]
            
            if processed_files is not None:
                update_fields.append('processed_files = ?')
                params.append(processed_files)
            
            if failed_files is not None:
                update_fields.append('failed_files = ?')
                params.append(failed_files)
            
            if error_message is not None:
                update_fields.append('error_message = ?')
                params.append(error_message)
            
            if status == 'processing':
                update_fields.append('started_at = CURRENT_TIMESTAMP')
            elif status in ['completed', 'failed']:
                update_fields.append('completed_at = CURRENT_TIMESTAMP')
            
            params.append(job_id)
            
            cursor.execute(f'''
                UPDATE upload_jobs 
                SET {', '.join(update_fields)}
                WHERE job_id = ?
            ''', params)
            
            conn.commit()
            logger.info(f"Updated job {job_id} status to {status}")
        finally:
            conn.close()

    def create_file_processing_status(self, job_id: str, filename: str, file_size: int, 
                                    file_hash: str) -> str:
        """Create individual file processing status"""
        status_id = f"status_{uuid.uuid4().hex[:8]}"
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT INTO file_processing_status 
                (status_id, job_id, filename, file_size, file_hash, status)
                VALUES (?, ?, ?, ?, ?, 'pending')
            ''', (status_id, job_id, filename, file_size, file_hash))
            conn.commit()
            return status_id
        finally:
            conn.close()

    def update_file_processing_status(self, status_id: str, status: str, 
                                    error_message: str = None, document_id: str = None, 
                                    chunks_created: int = None):
        """Update individual file processing status"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            update_fields = ['status = ?']
            params = [status]
            
            if error_message is not None:
                update_fields.append('error_message = ?')
                params.append(error_message)
            
            if document_id is not None:
                update_fields.append('document_id = ?')
                params.append(document_id)
            
            if chunks_created is not None:
                update_fields.append('chunks_created = ?')
                params.append(chunks_created)
            
            if status == 'processing':
                update_fields.append('processing_started_at = CURRENT_TIMESTAMP')
            elif status in ['completed', 'failed']:
                update_fields.append('processing_completed_at = CURRENT_TIMESTAMP')
            
            params.append(status_id)
            
            cursor.execute(f'''
                UPDATE file_processing_status 
                SET {', '.join(update_fields)}
                WHERE status_id = ?
            ''', params)
            
            conn.commit()
        finally:
            conn.close()

class EnhancedMeetingDocumentProcessor:
    """Enhanced Meeting Document Processor with Vector Database Support"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize the enhanced processor"""
        global llm, embedding_model, access_token
        self.llm = llm
        self.embedding_model = embedding_model
        self.access_token = access_token
        self.token_expiry = datetime.now() + timedelta(hours=1)
        
        # Text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
        )
        
        # Vector database
        self.vector_db = VectorDatabase()
        
        logger.info("Enhanced Meeting Document Processor initialized with OpenAI")
    
    def _detect_timeframe_from_query(self, query: str) -> Optional[str]:
        """Enhanced timeframe detection from natural language query"""
        query_lower = query.lower()
        
        # Comprehensive timeframe patterns with priority
        timeframe_patterns = [
            # Current periods (highest priority)
            (['current week', 'this week'], 'current_week'),
            (['current month', 'this month'], 'current_month'),
            (['current quarter', 'this quarter'], 'current_quarter'),
            (['current year', 'this year'], 'current_year'),
            
            # Last periods (high priority)
            (['last week', 'past week', 'previous week'], 'last_week'),
            (['last month', 'past month', 'previous month'], 'last_month'),
            (['last quarter', 'past quarter', 'previous quarter'], 'last_quarter'),
            (['last year', 'past year', 'previous year'], 'last_year'),
            
            # Specific day counts (medium priority)
            (['last 7 days', 'past 7 days', 'last seven days'], 'last_7_days'),
            (['last 14 days', 'past 14 days', 'last two weeks'], 'last_14_days'),
            (['last 30 days', 'past 30 days', 'last thirty days'], 'last_30_days'),
            (['last 60 days', 'past 60 days', 'last sixty days'], 'last_60_days'),
            (['last 90 days', 'past 90 days', 'last ninety days'], 'last_90_days'),
            
            # Extended periods (lower priority)
            (['last 3 months', 'past 3 months', 'last three months'], 'last_3_months'),
            (['last 6 months', 'past 6 months', 'last six months'], 'last_6_months'),
            (['last 12 months', 'past 12 months', 'last twelve months'], 'last_12_months'),
            
            # Recent periods (lowest priority)
            (['recent', 'recently', 'lately'], 'recent'),
        ]
        
        # Find the best match (earliest occurrence has highest priority)
        for patterns, timeframe in timeframe_patterns:
            for pattern in patterns:
                if pattern in query_lower:
                    return timeframe
        
        return None
    
    def _generate_date_based_summary(self, query: str, documents: List[Any], timeframe: str, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Generate intelligent date-based summary with chronological organization"""
        
        # Sort documents by date
        sorted_docs = sorted(documents, key=lambda x: x.date)
        
        # Group documents by date for better organization
        from collections import defaultdict
        date_groups = defaultdict(list)
        for doc in sorted_docs:
            date_key = doc.date.strftime('%Y-%m-%d')
            date_groups[date_key].append(doc)
        
        # Build comprehensive context from all documents
        context_parts = []
        document_summaries = []
        
        for date_key, docs in sorted(date_groups.items()):
            date_formatted = datetime.strptime(date_key, '%Y-%m-%d').strftime('%B %d, %Y')
            context_parts.append(f"\n=== {date_formatted} ===")
            
            for doc in docs:
                # Add document summary
                doc_summary = f"Document: {doc.filename}\n"
                if doc.content_summary:
                    doc_summary += f"Summary: {doc.content_summary}\n"
                if doc.main_topics:
                    doc_summary += f"Main Topics: {', '.join(doc.main_topics)}\n"
                if doc.participants:
                    doc_summary += f"Participants: {', '.join(doc.participants)}\n"
                if doc.future_actions:
                    doc_summary += f"Action Items: {', '.join(doc.future_actions)}\n"
                
                context_parts.append(doc_summary)
                document_summaries.append(doc_summary)
        
        # Create comprehensive context
        full_context = '\n'.join(context_parts)
        
        # Generate summary prompt based on query type
        timeframe_display = timeframe.replace('_', ' ').title()
        summary_prompt = f"""
        The user asked: "{query}"
        
        Based on the meeting documents from {timeframe_display}, please answer their question naturally and comprehensively. Focus on what they specifically asked for rather than forcing a predetermined structure.

        Meeting Documents Context:
        {full_context}

        Provide a direct, helpful answer that addresses exactly what the user wants to know about {timeframe_display}.
        
        IMPORTANT: When referencing information, always cite the specific document filename rather than document numbers or chunk references.
        """
        
        try:
            # Use class LLM instance
            messages = [
                SystemMessage(content="You are a helpful AI assistant that answers questions about meeting documents naturally. Focus on what the user specifically asked for rather than forcing a predetermined structure. Always cite document filenames rather than chunk numbers when referencing information."),
                HumanMessage(content=summary_prompt)
            ]
            
            response = self.llm.invoke(messages)
            summary_response = response.content.strip()
            
            # Add timeframe context to the response
            final_response = f"**Summary for {timeframe_display}** ({len(documents)} documents)\n\n{summary_response}"
            
            if include_context:
                return final_response, full_context
            else:
                return final_response
                
        except Exception as e:
            logger.error(f"Error generating date-based summary: {e}")
            # Enhanced fallback using stored content summaries
            fallback_parts = [f"**Summary for {timeframe_display}** ({len(documents)} documents)\n"]
            
            for date_key, docs in sorted(date_groups.items()):
                date_formatted = datetime.strptime(date_key, '%Y-%m-%d').strftime('%B %d, %Y')
                fallback_parts.append(f"\n**{date_formatted}:**")
                
                for doc in docs:
                    fallback_parts.append(f"\n• **{doc.filename}**")
                    
                    # Include actual content summary if available
                    if doc.content_summary:
                        fallback_parts.append(f"  {doc.content_summary}")
                    
                    # Include main topics if available  
                    if doc.main_topics:
                        topics_str = ', '.join(doc.main_topics[:3])  # First 3 topics
                        fallback_parts.append(f"  *Topics: {topics_str}*")
                    
                    # Include participants if available
                    if doc.participants:
                        participants_str = ', '.join(doc.participants[:3])  # First 3 participants
                        fallback_parts.append(f"  *Participants: {participants_str}*")
            
            fallback_summary = '\n'.join(fallback_parts)
            
            if include_context:
                return fallback_summary, full_context
            else:
                return fallback_summary
    def refresh_clients(self):
        """Refresh clients - simplified for OpenAI (no token refresh needed)"""
        try:
            global access_token, llm, embedding_model
            
            # Force reload environment variables
            load_dotenv(override=True)
            
            # Check API key
            current_api_key = os.getenv("OPENAI_API_KEY")
            if not current_api_key:
                raise ValueError("OPENAI_API_KEY not found in environment")
            
            logger.info("Refreshing LLM client with new API key")
            
            # For OpenAI, we don't need to refresh tokens since we use API keys
            # But we can recreate the clients if needed
            access_token = get_access_token()  # This just returns a dummy token
            llm = get_llm(access_token)
            embedding_model = get_embedding_model(access_token)
            
            self.llm = llm
            self.embedding_model = embedding_model
            self.access_token = access_token
            
            logger.info("OpenAI clients refreshed successfully")
            
        except Exception as e:
            logger.error(f"Failed to refresh clients: {e}")
            raise
    
    def extract_date_from_filename(self, filename: str) -> datetime:
        """Extract date from filename pattern with multiple formats"""
        patterns = [
            r'(\d{8})_(\d{6})',  # YYYYMMDD_HHMMSS
            r'(\d{8})',          # YYYYMMDD
            r'(\d{4}-\d{2}-\d{2})', # YYYY-MM-DD
            r'(\d{2}/\d{2}/\d{4})'  # MM/DD/YYYY
        ]
        
        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                try:
                    if '_' in match.group(0):
                        return datetime.strptime(match.group(0), "%Y%m%d_%H%M%S")
                    elif '-' in match.group(0):
                        return datetime.strptime(match.group(0), "%Y-%m-%d")
                    elif '/' in match.group(0):
                        return datetime.strptime(match.group(0), "%m/%d/%Y")
                    else:
                        return datetime.strptime(match.group(0), "%Y%m%d")
                except ValueError:
                    continue
        
        logger.warning(f"Could not extract date from filename: {filename}, using current date")
        return datetime.now()
    
    def read_document_content(self, file_path: str) -> str:
        """Read document content from file"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
            elif file_ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except ImportError:
                    logger.error("python-docx not installed. Cannot process .docx files.")
                    return ""
                    
            elif file_ext == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                except ImportError:
                    logger.error("PyPDF2 not installed. Cannot process .pdf files.")
                    return ""
            else:
                logger.error(f"Unsupported file format: {file_ext}")
                return ""
            
            return content if content.strip() else ""
                
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return ""
    
    def create_content_summary(self, content: str, max_length: int = 1500) -> str:
        """Create a condensed summary of the content"""
        try:
            summary_prompt = f"""
            Create a concise summary of this meeting document in 2-3 sentences (max {max_length} characters).
            Focus on the main purpose, key decisions, and outcomes.
            
            Content: {content[:2000]}...
            
            Summary:
            """
            
            messages = [
                SystemMessage(content="You are a meeting summarization expert. Create concise, informative summaries."),
                HumanMessage(content=summary_prompt)
            ]
            
            response = self.llm.invoke(messages)
            summary = response.content.strip()
            
            return summary[:max_length] if len(summary) > max_length else summary
            
        except Exception as e:
            logger.error(f"Error creating content summary: {e}")
            # Fallback to truncated content
            return content[:max_length] + "..." if len(content) > max_length else content
    
    def parse_document_content(self, content: str, filename: str, user_id: str, project_id: str = None, meeting_id: str = None) -> MeetingDocument:
        """Parse a meeting document and extract structured information"""
        doc_date = self.extract_date_from_filename(filename)
        document_id = f"{filename}_{doc_date.strftime('%Y%m%d_%H%M%S')}"
        
        parsing_prompt = f"""
        You are an expert document analyst. Analyze this meeting document and extract information in valid JSON format.
        
        Document content:
        {content[:4000]}{"..." if len(content) > 4000 else ""}
        
        Extract:
        1. "title": Clear, descriptive meeting title (2-8 words)
        2. "main_topics": Array of main topics discussed
        3. "past_events": Array of past events/completed items mentioned
        4. "future_actions": Array of upcoming actions/planned activities
        5. "participants": Array of participant names mentioned
        
        Return only valid JSON with no additional text.
        """
        
        try:
            messages = [
                SystemMessage(content="You are a document parsing assistant. Always return valid JSON format with no additional text."),
                HumanMessage(content=parsing_prompt)
            ]
            
            response = self.llm.invoke(messages)
            content_str = response.content.strip()
            
            # Clean JSON response
            if content_str.startswith('```json'):
                content_str = content_str[7:-3].strip()
            elif content_str.startswith('```'):
                content_str = content_str[3:-3].strip()
            
            parsed_data = json.loads(content_str)
            
            # Create content summary
            content_summary = self.create_content_summary(content)
            
            return MeetingDocument(
                document_id=document_id,
                filename=filename,
                date=doc_date,
                title=parsed_data.get('title', f'Meeting - {doc_date.strftime("%Y-%m-%d")}'),
                content=content,
                content_summary=content_summary,
                main_topics=parsed_data.get('main_topics', []),
                past_events=parsed_data.get('past_events', []),
                future_actions=parsed_data.get('future_actions', []),
                participants=parsed_data.get('participants', []),
                file_size=len(content),
                user_id=user_id,
                project_id=project_id,
                meeting_id=meeting_id
            )
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            return self._create_fallback_document(content, filename, doc_date, document_id, user_id, project_id, meeting_id)
    
    def _create_fallback_document(self, content: str, filename: str, doc_date: datetime, document_id: str, user_id: str, project_id: str = None, meeting_id: str = None) -> MeetingDocument:
        """Create a fallback document when parsing fails"""
        content_summary = self.create_content_summary(content)
        
        return MeetingDocument(
            document_id=document_id,
            filename=filename,
            date=doc_date,
            title=f"Meeting - {doc_date.strftime('%Y-%m-%d')}",
            content=content,
            content_summary=content_summary,
            main_topics=[],
            past_events=[],
            future_actions=[],
            participants=[],
            file_size=len(content),
            user_id=user_id,
            project_id=project_id,
            meeting_id=meeting_id
        )
    
    def extract_meeting_intelligence(self, content: str, max_tokens: int = 115000) -> Dict:
        """Extract comprehensive meeting intelligence using LLM"""
        try:
            # Estimate content tokens and handle large meetings
            estimated_tokens = len(content.split()) * 1.3  # Rough estimation
            
            if estimated_tokens > max_tokens:
                # For large meetings, process in windows
                return self._extract_intelligence_windowed(content, max_tokens)
            else:
                # Single pass for smaller meetings
                return self._extract_intelligence_single_pass(content)
                
        except Exception as e:
            logger.error(f"Error extracting meeting intelligence: {e}")
            return self._create_fallback_intelligence()
    
    def _extract_intelligence_single_pass(self, content: str) -> Dict:
        """Extract intelligence in single LLM call"""
        intelligence_prompt = f"""
        COMPREHENSIVE MEETING INTELLIGENCE EXTRACTION
        
        Meeting Content:
        {content}
        
        Extract ALL information and structure as JSON:
        
        {{
            "meeting_metadata": {{
                "meeting_type": "standup/planning/demo/review/discussion",
                "main_purpose": "brief description of meeting purpose",
                "duration_estimate": "estimated duration if mentioned",
                "meeting_effectiveness": "productive/average/ineffective"
            }},
            "participants": [
                {{
                    "name": "Full Name",
                    "role": "role if mentioned",
                    "contribution_summary": "what they primarily discussed",
                    "speaking_frequency": "high/medium/low",
                    "key_statements": ["important quotes"],
                    "questions_asked": ["questions they posed"],
                    "decisions_influenced": ["decisions they affected"],
                    "expertise_demonstrated": ["areas of knowledge shown"]
                }}
            ],
            "topics_discussed": [
                {{
                    "topic": "topic name",
                    "discussed_by": ["speaker1", "speaker2"],
                    "key_points": ["point1", "point2"],
                    "outcome": "decision/action/discussion",
                    "outcome_details": "specific outcome",
                    "importance_score": 0.8
                }}
            ],
            "decisions_made": [
                {{
                    "decision": "what was decided",
                    "decided_by": "who made decision",
                    "context": "why this decision",
                    "impact": "what this affects",
                    "implementation_notes": "how to implement"
                }}
            ],
            "action_items": [
                {{
                    "task": "what needs to be done",
                    "assigned_to": "person responsible",
                    "due_date": "if mentioned",
                    "priority": "high/medium/low",
                    "dependencies": ["other tasks"],
                    "context": "why this task is needed"
                }}
            ],
            "questions_and_answers": [
                {{
                    "question": "what was asked",
                    "asked_by": "who asked",
                    "answered_by": "who responded",
                    "answer": "response given",
                    "resolved": true
                }}
            ],
            "meeting_flow": {{
                "opening_discussion": "how meeting started",
                "main_discussion_points": ["point1", "point2"],
                "key_transitions": ["transition descriptions"],
                "closing_notes": "how meeting ended"
            }},
            "contextual_references": {{
                "previous_meetings_mentioned": ["meeting references"],
                "external_projects_mentioned": ["project names"],
                "systems_discussed": ["system names"],
                "deadlines_mentioned": ["deadline references"]
            }},
            "searchable_metadata": {{
                "all_names_mentioned": ["name1", "name2"],
                "technical_terms": ["term1", "term2"], 
                "project_names": ["project1", "project2"],
                "key_phrases": ["phrase1", "phrase2"],
                "sentiment_tone": "positive/neutral/negative",
                "urgency_level": "high/medium/low"
            }}
        }}
        
        Be extremely thorough. Extract every piece of information that could be useful for future queries.
        Include speaker attributions for all content. Preserve context and nuance.
        """
        
        try:
            messages = [
                SystemMessage(content="You are an expert meeting intelligence analyst. Always return valid JSON with comprehensive information extraction."),
                HumanMessage(content=intelligence_prompt)
            ]
            
            response = self.llm.invoke(messages)
            content_str = response.content.strip()
            
            # Clean JSON response
            if content_str.startswith('```json'):
                content_str = content_str[7:-3].strip()
            elif content_str.startswith('```'):
                content_str = content_str[3:-3].strip()
            
            intelligence_data = json.loads(content_str)
            logger.info("Successfully extracted meeting intelligence")
            return intelligence_data
            
        except Exception as e:
            logger.error(f"Error in single-pass intelligence extraction: {e}")
            return self._create_fallback_intelligence()
    
    def _extract_intelligence_windowed(self, content: str, max_tokens: int) -> Dict:
        """Extract intelligence from large meetings using windowed approach"""
        # For now, use truncated content for large meetings
        # TODO: Implement full windowed processing if needed
        truncated_content = content[:max_tokens * 3]  # Rough token-to-char conversion
        logger.warning(f"Content truncated for intelligence extraction: {len(content)} -> {len(truncated_content)} characters")
        return self._extract_intelligence_single_pass(truncated_content)
    
    def _create_fallback_intelligence(self) -> Dict:
        """Create fallback intelligence structure when extraction fails"""
        return {
            "meeting_metadata": {
                "meeting_type": "discussion",
                "main_purpose": "General meeting discussion",
                "duration_estimate": "unknown",
                "meeting_effectiveness": "average"
            },
            "participants": [],
            "topics_discussed": [],
            "decisions_made": [],
            "action_items": [],
            "questions_and_answers": [],
            "meeting_flow": {
                "opening_discussion": "Meeting discussion",
                "main_discussion_points": [],
                "key_transitions": [],
                "closing_notes": "Meeting concluded"
            },
            "contextual_references": {
                "previous_meetings_mentioned": [],
                "external_projects_mentioned": [],
                "systems_discussed": [],
                "deadlines_mentioned": []
            },
            "searchable_metadata": {
                "all_names_mentioned": [],
                "technical_terms": [],
                "project_names": [],
                "key_phrases": [],
                "sentiment_tone": "neutral",
                "urgency_level": "medium"
            }
        }
    
    def chunk_document(self, document: MeetingDocument, intelligence_data: Dict = None) -> List[DocumentChunk]:
        """Split document into intelligent chunks with enhanced metadata"""
        # Extract intelligence if not provided
        if intelligence_data is None:
            intelligence_data = self.extract_meeting_intelligence(document.content)
        
        # Create intelligent chunks with context preservation
        intelligent_chunks = self._create_intelligent_chunks(document, intelligence_data)
        
        document.chunk_count = len(intelligent_chunks)
        return intelligent_chunks
    
    def _create_intelligent_chunks(self, document: MeetingDocument, intelligence_data: Dict) -> List[DocumentChunk]:
        """Create intelligent chunks with rich metadata"""
        # Split content into base chunks
        base_chunks = self.text_splitter.split_text(document.content)
        
        document_chunks = []
        current_pos = 0
        
        for i, chunk_content in enumerate(base_chunks):
            chunk_id = f"{document.document_id}_chunk_{i}"
            
            # Find start position in original content
            start_char = document.content.find(chunk_content, current_pos)
            if start_char == -1:
                start_char = current_pos
            
            end_char = start_char + len(chunk_content)
            current_pos = end_char
            
            # Generate embedding for chunk
            try:
                if self.embedding_model is None:
                    logger.warning("Embedding model not available, using zero vector")
                    embedding_array = np.zeros(3072)
                else:
                    embedding = self.embedding_model.embed_query(chunk_content)
                    embedding_array = np.array(embedding)
            except Exception as e:
                logger.error(f"Error generating embedding for chunk {chunk_id}: {e}")
                embedding_array = np.zeros(3072)
            
            # Extract chunk-specific intelligence
            chunk_intelligence = self._extract_chunk_intelligence(chunk_content, intelligence_data, i, len(base_chunks))
            
            # Create enhanced chunk with metadata
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                document_id=document.document_id,
                filename=document.filename,
                chunk_index=i,
                content=chunk_content,
                start_char=start_char,
                end_char=end_char,
                embedding=embedding_array
            )
            
            # Add intelligence metadata
            chunk.enhanced_content = self._create_enhanced_content(chunk_content, chunk_intelligence)
            chunk.chunk_type = chunk_intelligence.get('chunk_type', 'discussion')
            chunk.speakers = json.dumps(chunk_intelligence.get('speakers', []))
            chunk.speaker_contributions = json.dumps(chunk_intelligence.get('speaker_contributions', {}))
            chunk.topics = json.dumps(chunk_intelligence.get('topics', []))
            chunk.decisions = json.dumps(chunk_intelligence.get('decisions', []))
            chunk.actions = json.dumps(chunk_intelligence.get('actions', []))
            chunk.questions = json.dumps(chunk_intelligence.get('questions', []))
            chunk.context_before = chunk_intelligence.get('context_before', '')
            chunk.context_after = chunk_intelligence.get('context_after', '')
            chunk.key_phrases = json.dumps(chunk_intelligence.get('key_phrases', []))
            chunk.importance_score = chunk_intelligence.get('importance_score', 0.5)
            
            document_chunks.append(chunk)
        
        return document_chunks
    
    def _extract_chunk_intelligence(self, chunk_content: str, meeting_intelligence: Dict, chunk_index: int, total_chunks: int) -> Dict:
        """Extract intelligence metadata for a specific chunk"""
        chunk_intel = {
            'chunk_type': 'discussion',
            'speakers': [],
            'speaker_contributions': {},
            'topics': [],
            'decisions': [],
            'actions': [],
            'questions': [],
            'context_before': '',
            'context_after': '',
            'key_phrases': [],
            'importance_score': 0.5
        }
        
        try:
            # Extract speakers mentioned in this chunk
            all_participants = meeting_intelligence.get('participants', [])
            chunk_speakers = []
            chunk_speaker_contributions = {}
            
            logger.info(f"Processing chunk {chunk_index}: Found {len(all_participants)} participants in intelligence data")
            if all_participants:
                logger.info(f"Participants: {[p.get('name', 'Unknown') for p in all_participants]}")
            
            for participant in all_participants:
                participant_name = participant.get('name', '')
                if participant_name:
                    # Enhanced name matching to handle different formats
                    name_parts = participant_name.split()
                    chunk_lower = chunk_content.lower()
                    
                    # Check multiple matching strategies
                    matched = False
                    
                    # Strategy 1: Exact name match
                    if participant_name.lower() in chunk_lower:
                        matched = True
                    
                    # Strategy 2: Last name, First name format (common in meeting transcripts)
                    elif len(name_parts) >= 2:
                        # Try "Last, First" format
                        last_first = f"{name_parts[-1]}, {' '.join(name_parts[:-1])}"
                        if last_first.lower() in chunk_lower:
                            matched = True
                        
                        # Try individual name parts
                        elif any(part.lower() in chunk_lower for part in name_parts if len(part) > 2):
                            matched = True
                    
                    # Strategy 3: Check if any significant name part appears
                    elif any(part.lower() in chunk_lower for part in name_parts if len(part) > 2):
                        matched = True
                    
                    if matched:
                        chunk_speakers.append(participant_name)
                        chunk_speaker_contributions[participant_name] = {
                            'contribution_summary': participant.get('contribution_summary', ''),
                            'speaking_frequency': participant.get('speaking_frequency', 'low')
                        }
                        logger.info(f"Found speaker {participant_name} in chunk {chunk_index}")
            
            if not chunk_speakers and all_participants:
                logger.info(f"No speakers matched in chunk {chunk_index}. First 200 chars: {chunk_content[:200]}")
            
            # Extract topics relevant to this chunk
            chunk_topics = []
            for topic in meeting_intelligence.get('topics_discussed', []):
                topic_name = topic.get('topic', '')
                if any(keyword.lower() in chunk_content.lower() for keyword in topic.get('key_points', [])):
                    chunk_topics.append(topic_name)
            
            # Extract decisions relevant to this chunk
            chunk_decisions = []
            for decision in meeting_intelligence.get('decisions_made', []):
                decision_text = decision.get('decision', '')
                if any(word in chunk_content.lower() for word in decision_text.lower().split()[:5]):
                    chunk_decisions.append(decision)
            
            # Extract action items relevant to this chunk
            chunk_actions = []
            for action in meeting_intelligence.get('action_items', []):
                action_text = action.get('task', '')
                if any(word in chunk_content.lower() for word in action_text.lower().split()[:5]):
                    chunk_actions.append(action)
            
            # Extract Q&A relevant to this chunk
            chunk_questions = []
            for qa in meeting_intelligence.get('questions_and_answers', []):
                question_text = qa.get('question', '')
                if any(word in chunk_content.lower() for word in question_text.lower().split()[:5]):
                    chunk_questions.append(qa)
            
            # Determine chunk type based on content
            chunk_type = 'discussion'
            if chunk_decisions:
                chunk_type = 'decision'
            elif chunk_actions:
                chunk_type = 'action_planning'
            elif chunk_questions:
                chunk_type = 'qa'
            
            # Calculate importance score
            importance_score = 0.5
            if chunk_decisions:
                importance_score += 0.3
            if chunk_actions:
                importance_score += 0.2
            if len(chunk_speakers) > 1:
                importance_score += 0.1
            importance_score = min(importance_score, 1.0)
            
            # Extract key phrases from chunk
            chunk_phrases = []
            searchable_metadata = meeting_intelligence.get('searchable_metadata', {})
            for phrase in searchable_metadata.get('key_phrases', []):
                if phrase.lower() in chunk_content.lower():
                    chunk_phrases.append(phrase)
            
            # Set context
            context_before = f"Chunk {chunk_index + 1} of {total_chunks}"
            context_after = f"Part of {meeting_intelligence.get('meeting_metadata', {}).get('main_purpose', 'meeting discussion')}"
            
            # Update chunk intelligence
            chunk_intel.update({
                'chunk_type': chunk_type,
                'speakers': chunk_speakers,
                'speaker_contributions': chunk_speaker_contributions,
                'topics': chunk_topics,
                'decisions': chunk_decisions,
                'actions': chunk_actions,
                'questions': chunk_questions,
                'context_before': context_before,
                'context_after': context_after,
                'key_phrases': chunk_phrases,
                'importance_score': importance_score
            })
            
        except Exception as e:
            logger.error(f"Error extracting chunk intelligence: {e}")
        
        return chunk_intel
    
    def _create_enhanced_content(self, original_content: str, chunk_intelligence: Dict) -> str:
        """Create enhanced content with context and speaker attribution"""
        enhanced_parts = []
        
        # Add context header
        context_before = chunk_intelligence.get('context_before', '')
        if context_before:
            enhanced_parts.append(f"[Context: {context_before}]")
        
        # Add speaker information
        speakers = chunk_intelligence.get('speakers', [])
        if speakers:
            enhanced_parts.append(f"[Speakers: {', '.join(speakers)}]")
        
        # Add chunk type
        chunk_type = chunk_intelligence.get('chunk_type', 'discussion')
        enhanced_parts.append(f"[Type: {chunk_type}]")
        
        # Add original content
        enhanced_parts.append(original_content)
        
        # Add summary of decisions/actions if present
        decisions = chunk_intelligence.get('decisions', [])
        if decisions:
            decision_summary = "; ".join([d.get('decision', '') for d in decisions[:2]])
            enhanced_parts.append(f"[Decisions: {decision_summary}]")
        
        actions = chunk_intelligence.get('actions', [])
        if actions:
            action_summary = "; ".join([a.get('task', '') for a in actions[:2]])
            enhanced_parts.append(f"[Actions: {action_summary}]")
        
        return "\n".join(enhanced_parts)
    
    def process_documents(self, document_folder: str) -> None:
        """Process all documents in a folder with chunking and vector storage"""
        folder_path = Path(document_folder)
        
        if not folder_path.exists():
            logger.error(f"Folder {document_folder} does not exist")
            return
        
        # Get supported files
        supported_extensions = ['.docx', '.txt', '.pdf']
        doc_files = []
        for ext in supported_extensions:
            doc_files.extend(folder_path.glob(f"*{ext}"))
        
        
        processed_count = 0
        for doc_file in doc_files:
            try:
                logger.info(f"Processing: {doc_file.name}")
                
                # Read document content
                content = self.read_document_content(str(doc_file))
                if not content.strip():
                    logger.warning(f"No content extracted from {doc_file.name}")
                    continue
                
                # Parse document
                meeting_doc = self.parse_document_content(content, doc_file.name)
                
                # Create chunks with embeddings
                logger.info(f"Creating chunks for {doc_file.name}")
                chunks = self.chunk_document(meeting_doc)
                
                # Store in vector database
                self.vector_db.add_document(meeting_doc, chunks)
                
                processed_count += 1
                logger.info(f"Successfully processed: {doc_file.name} ({len(chunks)} chunks)")
                
            except Exception as e:
                logger.error(f"Error processing {doc_file.name}: {e}")
        
        # Save vector index
        if processed_count > 0:
            self.vector_db.save_index()
            logger.info(f"Successfully processed {processed_count} documents")
    
    def hybrid_search(self, query: str, user_id: str, project_id: str = None, meeting_id: str = None, folder_path: str = None, top_k: int = 15, semantic_weight: float = 0.7) -> List[DocumentChunk]:
        """Perform hybrid search combining semantic and keyword search"""
        
        # Extract keywords from query
        keywords = [word.lower().strip() for word in query.split() if len(word) > 2]
        
        # Semantic search
        try:
            if self.embedding_model is None:
                logger.warning("Embedding model not available, skipping semantic search")
                semantic_results = []
            else:
                query_embedding = np.array(self.embedding_model.embed_query(query))
                if folder_path:
                    # Filter semantic search by folder
                    semantic_results = self.vector_db.search_similar_chunks_by_folder(query_embedding, user_id, folder_path, top_k * 2)
                else:
                    semantic_results = self.vector_db.search_similar_chunks(query_embedding, top_k * 2)
        except Exception as e:
            logger.error(f"Error in semantic search: {e}")
            semantic_results = []
        
        # Keyword search with user context
        if folder_path:
            keyword_chunk_ids = self.vector_db.keyword_search_chunks_by_folder(keywords, user_id, folder_path, top_k)
        else:
            keyword_chunk_ids = self.vector_db.keyword_search_chunks_by_user(keywords, user_id, project_id, meeting_id, top_k)
        
        # Combine and score results
        chunk_scores = defaultdict(float)
        
        # Add semantic scores
        for chunk_id, similarity in semantic_results:
            chunk_scores[chunk_id] += similarity * semantic_weight
        
        # Add keyword scores
        keyword_weight = 1.0 - semantic_weight
        for chunk_id in keyword_chunk_ids:
            chunk_scores[chunk_id] += keyword_weight * 0.5  # Base keyword score
        
        # Get top chunks
        top_chunk_ids = sorted(chunk_scores.keys(), key=lambda x: chunk_scores[x], reverse=True)[:top_k]
        
        # Retrieve chunk details
        return self.vector_db.get_chunks_by_ids(top_chunk_ids)
    
    def detect_summary_query(self, query: str) -> bool:
        """Detect if the query is asking for meeting summaries"""
        summary_keywords = [
            'summarize', 'summary', 'summaries', 'overview', 'brief', 
            'recap', 'highlights', 'key points', 'main points',
            'all meetings', 'all documents', 'overall', 'across all',
            'consolidate', 'aggregate', 'compile', 'comprehensive',
            'meetings summary', 'meeting summaries', 'summarize meetings',
            'summarize the meetings', 'summary of meetings', 'summary of all'
        ]
        
        query_lower = query.lower()
        for keyword in summary_keywords:
            if keyword in query_lower:
                return True
        return False

    def answer_query_with_intelligence(self, query: str, user_id: str, document_ids: List[str] = None, 
                                     project_id: str = None, meeting_id: str = None, meeting_ids: List[str] = None, 
                                     date_filters: List[str] = None, folder_path: str = None, 
                                     context_limit: int = 10, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Answer user query using enhanced intelligence-aware search and context reconstruction"""
        
        try:
            # Generate query embedding
            if self.embedding_model is None:
                logger.error("Embedding model not available")
                return "Sorry, the system is not properly configured for queries.", ""
            
            query_embedding = self.embedding_model.embed_query(query)
            query_vector = np.array(query_embedding)
            
            # Analyze query to determine filters
            search_filters = self._analyze_query_for_filters(query)
            
            # Apply user context filters
            if meeting_ids:
                search_filters['meeting_ids'] = meeting_ids
            if project_id:
                search_filters['project_id'] = project_id
            if date_filters:
                search_filters['date_filters'] = date_filters
            
            # Perform enhanced search with metadata filtering
            enhanced_results = self.vector_db.enhanced_search_with_metadata(
                query_vector, user_id, search_filters, top_k=context_limit
            )
            
            if not enhanced_results:
                return "I couldn't find any relevant information for your query. Please try rephrasing or check if you have uploaded meeting documents.", ""
            
            # Generate context-aware response
            response, context = self._generate_intelligence_response(query, enhanced_results, user_id)
            
            if include_context:
                return response, context
            else:
                return response
                
        except Exception as e:
            logger.error(f"Error in intelligence-aware query processing: {e}")
            return f"I encountered an error while processing your query: {str(e)}", ""
    
    def _analyze_query_for_filters(self, query: str) -> Dict:
        """Analyze query to determine appropriate metadata filters"""
        filters = {}
        query_lower = query.lower()
        
        # Detect speaker-specific queries
        speaker_patterns = ['what did', 'what said', 'who said', 'mentioned by', 'according to']
        if any(pattern in query_lower for pattern in speaker_patterns):
            # Try to extract speaker names from query
            potential_speakers = []
            words = query.split()
            
            # Enhanced pattern matching for speaker extraction
            for i, word in enumerate(words):
                word_lower = word.lower()
                if word_lower in ['what', 'who'] and i + 1 < len(words):
                    next_word = words[i + 1].lower()
                    if next_word in ['did', 'said']:
                        # Look for the speaker name after "what did" or "who said"
                        if i + 2 < len(words):
                            speaker_name = words[i + 2].strip('.,?!').title()
                            potential_speakers.append(speaker_name)
                
                # Also check for patterns like "sandeep said" or "according to john"
                if word_lower in ['said', 'mentioned'] and i > 0:
                    speaker_name = words[i - 1].strip('.,?!').title()
                    potential_speakers.append(speaker_name)
                elif word_lower == 'to' and i > 0 and words[i - 1].lower() == 'according':
                    if i + 1 < len(words):
                        speaker_name = words[i + 1].strip('.,?!').title()
                        potential_speakers.append(speaker_name)
            
            # Remove duplicates and common words
            common_words = {'the', 'in', 'at', 'on', 'for', 'with', 'by', 'from', 'meeting', 'discussion'}
            potential_speakers = list(set([s for s in potential_speakers if s.lower() not in common_words]))
            
            if potential_speakers:
                filters['speakers'] = potential_speakers
                logger.info(f"Extracted speakers from query '{query}': {potential_speakers}")
        
        # Detect decision-focused queries
        decision_patterns = ['decision', 'decided', 'conclusion', 'resolution', 'agreed']
        if any(pattern in query_lower for pattern in decision_patterns):
            filters['has_decisions'] = True
            filters['chunk_type'] = ['decision', 'discussion']
        
        # Detect action item queries
        action_patterns = ['action', 'task', 'todo', 'follow up', 'next steps', 'assigned']
        if any(pattern in query_lower for pattern in action_patterns):
            filters['has_actions'] = True
            filters['chunk_type'] = ['action_planning', 'discussion']
        
        # Detect high-importance queries
        importance_patterns = ['important', 'critical', 'urgent', 'key', 'major']
        if any(pattern in query_lower for pattern in importance_patterns):
            filters['min_importance'] = 0.7
        
        return filters
    
    def _generate_intelligence_response(self, query: str, enhanced_results: List[Dict], user_id: str) -> Tuple[str, str]:
        """Generate response using enhanced intelligence data"""
        
        # Prepare comprehensive context from enhanced results
        context_parts = []
        speaker_contributions = {}
        decisions_made = []
        actions_identified = []
        meeting_contexts = []
        
        for result in enhanced_results:
            chunk = result['chunk']
            context = result['context']
            metadata = result['metadata']
            
            # Collect context information
            context_part = f"**From {context['meeting_context'].get('meeting_title', 'Meeting')} ({context['meeting_context'].get('meeting_date', 'Unknown date')})**\n"
            context_part += f"Content: {chunk.enhanced_content or chunk.content}\n"
            
            if metadata['speakers']:
                context_part += f"Speakers: {', '.join(metadata['speakers'])}\n"
            
            if metadata['decisions']:
                context_part += f"Decisions: {'; '.join([d.get('decision', '') for d in metadata['decisions'][:2]])}\n"
                decisions_made.extend(metadata['decisions'])
            
            if metadata['actions']:
                context_part += f"Actions: {'; '.join([a.get('task', '') for a in metadata['actions'][:2]])}\n"
                actions_identified.extend(metadata['actions'])
            
            context_part += f"Relevance Score: {result['relevance_score']:.3f}\n\n"
            context_parts.append(context_part)
            
            # Collect speaker data
            for speaker in metadata['speakers']:
                if speaker not in speaker_contributions:
                    speaker_contributions[speaker] = []
                speaker_contributions[speaker].append({
                    'content': chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content,
                    'meeting': context['meeting_context'].get('meeting_title', 'Unknown'),
                    'date': context['meeting_context'].get('meeting_date', 'Unknown')
                })
            
            # Collect meeting context
            meeting_context = context['meeting_context']
            if meeting_context and meeting_context not in meeting_contexts:
                meeting_contexts.append(meeting_context)
        
        # Create comprehensive context string
        full_context = "COMPREHENSIVE MEETING INTELLIGENCE:\n\n" + "\n".join(context_parts)
        
        # Generate enhanced response prompt
        response_prompt = f"""
        Based on the comprehensive meeting intelligence provided, answer the user's query with complete accuracy and context.
        
        USER QUERY: "{query}"
        
        MEETING INTELLIGENCE CONTEXT:
        {full_context}
        
        SPEAKER ANALYSIS:
        {self._format_speaker_analysis(speaker_contributions)}
        
        DECISIONS SUMMARY:
        {self._format_decisions_summary(decisions_made)}
        
        ACTION ITEMS SUMMARY:
        {self._format_actions_summary(actions_identified)}
        
        INSTRUCTIONS:
        1. Answer the user's query directly and comprehensively
        2. Include specific speaker attributions when relevant
        3. Reference specific meetings and dates
        4. Include related decisions and action items
        5. Maintain chronological context when applicable
        6. Use the enhanced context to provide complete information
        
        Provide a detailed, accurate response that leverages all the intelligence available.
        """
        
        try:
            messages = [
                SystemMessage(content="You are an expert meeting intelligence assistant. Provide comprehensive, accurate responses using all available context and speaker attributions."),
                HumanMessage(content=response_prompt)
            ]
            
            response = self.llm.invoke(messages)
            return response.content.strip(), full_context
            
        except Exception as e:
            logger.error(f"Error generating intelligence response: {e}")
            return f"I found relevant information but encountered an error generating the response: {str(e)}", full_context
    
    def _format_speaker_analysis(self, speaker_contributions: Dict) -> str:
        """Format speaker contribution analysis"""
        if not speaker_contributions:
            return "No specific speaker contributions identified."
        
        analysis = []
        for speaker, contributions in speaker_contributions.items():
            contrib_summary = f"{speaker}: {len(contributions)} contributions across {len(set(c['meeting'] for c in contributions))} meetings"
            analysis.append(contrib_summary)
        
        return "\n".join(analysis)
    
    def _format_decisions_summary(self, decisions: List[Dict]) -> str:
        """Format decisions summary"""
        if not decisions:
            return "No specific decisions identified."
        
        summary = []
        for decision in decisions[:5]:  # Limit to top 5
            dec_text = decision.get('decision', 'Unknown decision')
            decided_by = decision.get('decided_by', 'Unknown')
            summary.append(f"- {dec_text} (by {decided_by})")
        
        return "\n".join(summary)
    
    def _format_actions_summary(self, actions: List[Dict]) -> str:
        """Format action items summary"""
        if not actions:
            return "No specific action items identified."
        
        summary = []
        for action in actions[:5]:  # Limit to top 5
            task = action.get('task', 'Unknown task')
            assigned_to = action.get('assigned_to', 'Unknown')
            summary.append(f"- {task} (assigned to {assigned_to})")
        
        return "\n".join(summary)

    def answer_query(self, query: str, user_id: str, document_ids: List[str] = None, project_id: str = None, meeting_id: str = None, meeting_ids: List[str] = None, date_filters: List[str] = None, folder_path: str = None, context_limit: int = 10, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Answer user query using hybrid search and intelligent context selection"""
        
        # Check if this is a summary query and no specific documents are selected
        is_summary_query = self.detect_summary_query(query)
        
        # Check if this is a comprehensive project summary query
        is_project_summary_query = self.detect_project_summary_query(query)
        
        # Handle comprehensive project summary requests
        if is_project_summary_query and not document_ids:
            logger.info("Detected comprehensive project summary query")
            return self._generate_comprehensive_project_summary(query, user_id, project_id, include_context)
        
        # Handle enhanced @ mention filters
        if meeting_ids:
            logger.info(f"Enhanced meeting filters: {meeting_ids}")
            # Support for multiple meeting IDs - use all provided meeting IDs
            if meeting_ids and not meeting_id:
                # Use all meeting IDs for document retrieval
                meeting_id = meeting_ids  # Pass the full list instead of just first one
        
        if date_filters:
            logger.info(f"Date filters: {date_filters}")
            # Apply explicit date filtering using existing timeframe logic
            date_filtered_docs = []
            for date_filter in date_filters:
                try:
                    # Use existing timeframe detection and date filtering
                    timeframe_docs = self.get_documents_by_timeframe(date_filter, user_id)
                    # Extract document IDs from MeetingDocument objects
                    filtered_doc_ids = [doc.document_id for doc in timeframe_docs]
                    date_filtered_docs.extend(filtered_doc_ids)
                    logger.info(f"Date filter '{date_filter}' matched {len(filtered_doc_ids)} documents")
                except Exception as e:
                    logger.warning(f"Error applying date filter '{date_filter}': {e}")
            
            # Remove duplicates and use date-filtered documents
            if date_filtered_docs:
                date_filtered_docs = list(set(date_filtered_docs))
                if not document_ids:
                    document_ids = date_filtered_docs
                else:
                    # Intersect with existing document filters
                    document_ids = list(set(document_ids) & set(date_filtered_docs))
                logger.info(f"Applied date filters, now using {len(document_ids)} documents")
        
        # If no specific documents are provided, include all user documents
        if not document_ids:
            logger.info("No specific documents provided, including all user documents")
            if folder_path:
                # Use folder-based filtering ONLY - don't mix with project/meeting filters
                document_ids = self.vector_db.get_user_documents_by_folder(user_id, folder_path)
            else:
                # Use standard scope-based filtering
                document_ids = self.vector_db.get_user_documents_by_scope(user_id, project_id, meeting_id)
                logger.info(f"Including {len(document_ids)} documents for user {user_id}")
        
        # Enhanced intelligent timeframe detection
        detected_timeframe = self._detect_timeframe_from_query(query)
        
        # Get relevant documents by timeframe if specified
        if detected_timeframe:
            logger.info(f"Detected timeframe: {detected_timeframe}")
            timeframe_docs = self.vector_db.get_documents_by_timeframe(detected_timeframe, user_id)
            if not timeframe_docs:
                error_msg = f"I don't have any meeting documents from the {detected_timeframe.replace('_', ' ')} timeframe."
                return (error_msg, "") if include_context else error_msg
            
            # Filter document_ids to only include documents from the detected timeframe
            if not document_ids:
                document_ids = [doc.document_id for doc in timeframe_docs]
                logger.info(f"Using {len(document_ids)} documents from {detected_timeframe} timeframe")
            else:
                # Intersect with timeframe documents
                timeframe_doc_ids = {doc.document_id for doc in timeframe_docs}
                document_ids = [doc_id for doc_id in document_ids if doc_id in timeframe_doc_ids]
                logger.info(f"Filtered to {len(document_ids)} documents in {detected_timeframe} timeframe")
            
            # Generate enhanced summary if this is a summary query with date context
            if is_summary_query and len(timeframe_docs) > 1:
                return self._generate_date_based_summary(query, timeframe_docs, detected_timeframe, include_context)
        
        # Perform hybrid search
        relevant_chunks = self.hybrid_search(query, user_id, project_id, meeting_id, folder_path, top_k=context_limit * 3)
        
        # Filter chunks by document IDs if specified
        if document_ids:  # Temporarily disabled @ feature
            original_count = len(relevant_chunks)
            relevant_chunks = [chunk for chunk in relevant_chunks if chunk.document_id in document_ids]
            if not relevant_chunks:
                error_msg = "I don't have any relevant information in the specified documents for your question."
                return (error_msg, "") if include_context else error_msg
        elif not relevant_chunks:
            error_msg = "I don't have any relevant meeting documents to answer your question."
            return (error_msg, "") if include_context else error_msg
        
        # Group chunks by document and select best representatives
        document_chunks = defaultdict(list)
        for chunk in relevant_chunks:
            document_chunks[chunk.document_id].append(chunk)
        
        # Select best chunks from each document (max 2 per document)
        selected_chunks = []
        for doc_id, chunks in document_chunks.items():
            # Sort chunks by position to maintain context
            chunks.sort(key=lambda x: x.chunk_index)
            selected_chunks.extend(chunks[:2])  # Max 2 chunks per document
        
        # Limit total chunks
        selected_chunks = selected_chunks[:context_limit]
        
        # Build context without chunk references
        context_parts = []
        current_doc = None
        
        for chunk in selected_chunks:
            if chunk.document_id != current_doc:
                # Add document header when switching documents
                if current_doc is not None:
                    context_parts.append("\n" + "="*60 + "\n")
                
                context_parts.append(f"Document: {chunk.filename}")
                current_doc = chunk.document_id
            
            # Add content without chunk numbering
            context_parts.append(chunk.content)
        
        context = "\n".join(context_parts)
        
        # Generate answer using OpenAI GPT-4o
        answer_prompt = f"""
User Question: {query}

Meeting Document Context:
{context}

Please answer the user's question naturally and directly based on the meeting documents. Focus on what they specifically asked for, using relevant details from the documents. Don't follow a rigid format - just provide a helpful, conversational response that addresses their actual question.

IMPORTANT: When referencing information from the documents, always cite the document filename (e.g., "Document_Fulfillment_AIML-20250714_153021-Meeting_Recording.docx") rather than chunk numbers. This helps users know which specific document the information comes from.
"""
        
        try:
            messages = [
                SystemMessage(content="You are a helpful AI assistant that answers questions about meeting documents naturally and conversationally. Answer exactly what the user asks for without forcing predetermined structures. Always cite document filenames rather than chunk numbers when referencing information."),
                HumanMessage(content=answer_prompt)
            ]
            
            response = self.llm.invoke(messages)
            return (response.content, context) if include_context else response.content
            
        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            # Try refreshing clients and retry
            try:
                self.refresh_clients()
                messages = [
                    SystemMessage(content="You are a helpful AI assistant that answers questions about meeting documents naturally and conversationally. Answer exactly what the user asks for without forcing predetermined structures. Always cite document filenames rather than chunk numbers when referencing information."),
                    HumanMessage(content=answer_prompt)
                ]
                response = self.llm.invoke(messages)
                return (response.content, context) if include_context else response.content
            except Exception as retry_error:
                error_msg = f"I encountered an error while processing your question. Please try again later. Error details: {str(retry_error)}"
                return (error_msg, "") if include_context else error_msg
    
    def generate_follow_up_questions(self, user_query: str, ai_response: str, context: str) -> List[str]:
        """Generate follow-up questions based on the user query, AI response, and document context"""
        try:
            follow_up_prompt = f"""
Based on the user's question, the AI response provided, and the meeting document context, generate 4-5 relevant follow-up questions that the user might want to ask next.

User's Question: {user_query}

AI Response: {ai_response}

Meeting Context: {context[:2000]}...

Generate follow-up questions that:
- Build upon the current conversation topic
- Explore related aspects mentioned in the documents
- Ask for deeper details about key points
- Inquire about timelines, next steps, or implications
- Connect to other relevant topics in the meetings

Return exactly 4-5 questions, each on a new line, without numbers or bullet points. Make them natural and conversational.
"""

            messages = [
                SystemMessage(content="You generate natural, conversational follow-up questions that help users explore their meeting documents further."),
                HumanMessage(content=follow_up_prompt)
            ]
            
            response = self.llm.invoke(messages)
            
            # Parse the response into individual questions
            questions = [q.strip() for q in response.content.split('\n') if q.strip() and len(q.strip()) > 10]
            
            # Ensure we have 4-5 questions
            if len(questions) > 5:
                questions = questions[:5]
            elif len(questions) < 4:
                # Add generic questions if we don't have enough
                generic_questions = [
                    "What were the key decisions made in these meetings?",
                    "What are the next steps mentioned?",
                    "Are there any deadlines or milestones discussed?"
                ]
                questions.extend(generic_questions[:5-len(questions)])
            
            return questions[:5]
            
        except Exception as e:
            logger.error(f"Error generating follow-up questions: {e}")
            # Return default follow-up questions
            return [
                "What were the main decisions made in the meetings?",
                "What are the upcoming deadlines or milestones?",
                "Are there any action items assigned?"
            ]
    
    def get_meeting_statistics(self) -> Dict[str, Any]:
        """Get simplified statistics about processed meetings"""
        try:
            conn = sqlite3.connect(self.vector_db.db_path)
            cursor = conn.cursor()
            
            # Get document counts
            cursor.execute("SELECT COUNT(*) FROM documents")
            total_docs = cursor.fetchone()[0]
            
            if total_docs == 0:
                return {"error": "No documents processed"}
            
            # Get date range
            cursor.execute("SELECT MIN(date), MAX(date) FROM documents")
            date_range = cursor.fetchone()
            
            # Get chunk statistics
            cursor.execute("SELECT COUNT(*), AVG(LENGTH(content)) FROM chunks")
            chunk_stats = cursor.fetchone()
            
            # Get monthly distribution (optional - can keep for internal tracking)
            cursor.execute("SELECT strftime('%Y-%m', date) as month, COUNT(*) FROM documents GROUP BY month ORDER BY month")
            monthly_counts = dict(cursor.fetchall())
            
            conn.close()
            
            stats = {
                "total_meetings": total_docs,
                "total_chunks": chunk_stats[0] if chunk_stats[0] else 0,
                "average_chunk_length": int(chunk_stats[1]) if chunk_stats[1] else 0,
                "vector_index_size": self.vector_db.index.ntotal if self.vector_db.index else 0,
                "date_range": {
                    "earliest": date_range[0] if date_range[0] else "N/A",
                    "latest": date_range[1] if date_range[1] else "N/A"
                },
                "meetings_per_month": monthly_counts
            }
            
            return stats
        except Exception as e:
            logger.error(f"Error generating statistics: {e}")
            return {"error": f"Failed to generate statistics: {e}"}

    def detect_project_summary_query(self, query: str) -> bool:
        """Detect if the query is asking for a comprehensive project summary"""
        project_summary_keywords = [
            'project summary', 'project summaries', 'summarize project', 'summarize the project',
            'summary of project', 'summary of all files', 'all files summary', 'comprehensive summary',
            'summarize all meetings', 'all meetings summary', 'overall project', 'entire project',
            'project overview', 'complete summary', 'full summary', 'all documents summary',
            'project recap', 'project highlights', 'all files in project', 'everything in project'
        ]
        
        query_lower = query.lower()
        for keyword in project_summary_keywords:
            if keyword in query_lower:
                return True
        return False

    def _generate_comprehensive_project_summary(self, query: str, user_id: str, project_id: str = None, include_context: bool = False) -> Union[str, Tuple[str, str]]:
        """Generate flexible comprehensive answer that processes ALL files based on user query"""
        try:
            logger.info(f"Generating user-centric comprehensive answer for user {user_id}, project {project_id}")
            
            # Get all documents in the project
            if project_id:
                documents = self.vector_db.get_project_documents(project_id, user_id)
            else:
                documents = self.vector_db.get_all_documents(user_id)
            
            if not documents:
                error_msg = "No documents found in the project to analyze."
                return (error_msg, "") if include_context else error_msg
            
            total_files = len(documents)
            logger.info(f"Found {total_files} files to process for user query: '{query}'")
            
            # Use single flexible processing approach
            return self._generate_flexible_comprehensive_answer(documents, query, total_files, include_context)
                
        except Exception as e:
            logger.error(f"Error generating comprehensive answer: {e}")
            error_msg = f"I encountered an error processing your question across all project files: {str(e)}"
            return (error_msg, "") if include_context else error_msg

    def _generate_flexible_comprehensive_answer(self, documents: List[Any], query: str, total_files: int, include_context: bool) -> Union[str, Tuple[str, str]]:
        """Single flexible function that processes ALL files based on user query intent"""
        try:
            logger.info(f"Processing {total_files} files with flexible approach for query: '{query}'")
            
            # Smart content selection based on file count
            if total_files <= 20:
                # Small projects: Use detailed content from all files
                content_chunks = self._get_detailed_content_from_all_files(documents)
                processing_note = f"Analyzed all {total_files} files individually"
            elif total_files <= 50:
                # Medium projects: Use smart sampling + summaries
                content_chunks = self._get_smart_sampled_content(documents, query)
                processing_note = f"Analyzed all {total_files} files using smart sampling"
            else:
                # Large projects: Use summarized content + key excerpts
                content_chunks = self._get_summarized_content_with_excerpts(documents, query)
                processing_note = f"Analyzed all {total_files} files using intelligent summarization"
            
            if not content_chunks:
                error_msg = "Unable to extract relevant content from the documents."
                return (error_msg, "") if include_context else error_msg
            
            # Build context for transparency
            context = f"Processing Strategy: {processing_note}\n" + "="*60 + "\n"
            context += "\n".join([f"File {i+1}: {chunk['filename']}" for i, chunk in enumerate(content_chunks[:10])])
            if len(content_chunks) > 10:
                context += f"\n... and {len(content_chunks) - 10} more files"
            
            # Generate flexible, user-centric response
            flexible_prompt = f"""
User Question: {query}

Document Content from {total_files} files:
{self._format_content_for_analysis(content_chunks)}

Please answer the user's question naturally and thoroughly using information from all the files. Focus on exactly what they asked for without forcing any predetermined structure.

IMPORTANT: When referencing information, always cite the specific document filename rather than document numbers or chunk references. This helps users identify the source document.
"""

            messages = [
                SystemMessage(content=f"You are an expert analyst with access to {total_files} meeting documents. Answer user questions naturally and comprehensively based on all available information. Always cite document filenames rather than document numbers when referencing information."),
                HumanMessage(content=flexible_prompt)
            ]
            
            response = self.llm.invoke(messages)
            
            # Add transparency note
            final_response = f"*Based on analysis of all {total_files} files in your project*\n\n{response.content}"
            
            return (final_response, context) if include_context else final_response
            
        except Exception as e:
            logger.error(f"Error in flexible comprehensive processing: {e}")
            error_msg = f"Error analyzing {total_files} files: {str(e)}"
            return (error_msg, "") if include_context else error_msg

    def _get_detailed_content_from_all_files(self, documents: List[Any]) -> List[Dict[str, str]]:
        """Get detailed content from all files for small projects (≤15 files)"""
        content_chunks = []
        for doc in documents:
            doc_info = self.vector_db.get_document_metadata(doc['document_id'])
            if doc_info:
                content_chunks.append({
                    'filename': doc_info['filename'],
                    'date': doc_info['date'][:10],
                    'content': doc_info['content_summary'],
                    'topics': ', '.join(doc_info['main_topics'][:3]),  # Top 3 topics
                    'participants': ', '.join(doc_info['participants'][:5])  # Top 5 participants
                })
        return content_chunks

    def _get_smart_sampled_content(self, documents: List[Any], query: str) -> List[Dict[str, str]]:
        """Get smart sampled content for medium projects (16-50 files)"""
        # For medium projects, get summaries from all files but prioritize based on relevance
        all_content = self._get_detailed_content_from_all_files(documents)
        
        # Simple relevance scoring based on query keywords
        query_words = set(query.lower().split())
        
        for chunk in all_content:
            # Score based on how many query words appear in content
            content_words = set(chunk['content'].lower().split())
            chunk['relevance_score'] = len(query_words.intersection(content_words))
        
        # Sort by relevance but keep all files
        all_content.sort(key=lambda x: x['relevance_score'], reverse=True)
        return all_content

    def _get_summarized_content_with_excerpts(self, documents: List[Any], query: str) -> List[Dict[str, str]]:
        """Get summarized content with key excerpts for large projects (50+ files)"""
        # For large projects, group by time periods and get summaries
        content_chunks = []
        
        # Group documents by month for better organization
        monthly_groups = {}
        for doc in documents:
            doc_info = self.vector_db.get_document_metadata(doc['document_id'])
            if doc_info:
                try:
                    import datetime
                    doc_date = datetime.datetime.fromisoformat(doc_info['date'].replace('Z', '+00:00'))
                    month_key = doc_date.strftime('%Y-%m')
                    if month_key not in monthly_groups:
                        monthly_groups[month_key] = []
                    monthly_groups[month_key].append(doc_info)
                except:
                    if 'unknown' not in monthly_groups:
                        monthly_groups['unknown'] = []
                    monthly_groups['unknown'].append(doc_info)
        
        # Create summaries for each time period
        for period, docs in monthly_groups.items():
            period_summary = f"Period {period} ({len(docs)} files): "
            topics = set()
            participants = set()
            
            for doc in docs:
                topics.update(doc['main_topics'][:2])  # Top 2 topics per doc
                participants.update(doc['participants'][:3])  # Top 3 participants per doc
            
            period_summary += f"Main topics: {', '.join(list(topics)[:5])}. "
            period_summary += f"Key participants: {', '.join(list(participants)[:8])}."
            
            content_chunks.append({
                'filename': f"{len(docs)} files from {period}",
                'date': period,
                'content': period_summary,
                'topics': ', '.join(list(topics)[:5]),
                'participants': ', '.join(list(participants)[:8])
            })
        
        return content_chunks

    def _format_content_for_analysis(self, content_chunks: List[Dict[str, str]]) -> str:
        """Format content chunks for AI analysis"""
        formatted_content = []
        
        for i, chunk in enumerate(content_chunks, 1):
            formatted_content.append(f"""
Document {i}: {chunk['filename']} ({chunk['date']})
Content: {chunk['content']}
Key Topics: {chunk['topics']}
Participants: {chunk['participants']}
{'='*60}""")
        
        return "\n".join(formatted_content)

    def process_file_async(self, file_path: str, filename: str, user_id: str,
                          project_id: str = None, meeting_id: str = None,
                          job_id: str = None, status_id: str = None) -> Dict:
        """Process a single file asynchronously"""
        try:
            # Update status
            if status_id:
                self.vector_db.update_file_processing_status(status_id, 'processing')
            
            # Calculate file hash
            file_hash = self.vector_db.calculate_file_hash(file_path)
            file_size = os.path.getsize(file_path)
            
            # Check for duplicates again (safety check)
            duplicate_info = self.vector_db.is_file_duplicate(file_hash, filename, user_id)
            if duplicate_info:
                if status_id:
                    self.vector_db.update_file_processing_status(status_id, 'skipped', 
                                                               f"Duplicate of {duplicate_info['original_filename']}")
                return {
                    'success': False,
                    'error': f'Duplicate file: {duplicate_info["original_filename"]}',
                    'is_duplicate': True,
                    'duplicate_info': duplicate_info
                }
            
            # Extract content
            content = self.read_document_content(file_path)
            if not content.strip():
                if status_id:
                    self.vector_db.update_file_processing_status(status_id, 'failed', 'Empty file content')
                return {'success': False, 'error': 'Empty file content'}
            
            # Store file metadata
            document_id = self.vector_db.store_document_metadata(
                filename, content, user_id, project_id, meeting_id
            )
            
            # Store file hash for deduplication
            hash_id = self.vector_db.store_file_hash(
                file_hash, filename, filename, file_size, user_id,
                project_id, meeting_id, document_id
            )
            
            # Create MeetingDocument object for processing
            meeting_doc = MeetingDocument(
                document_id=document_id,
                filename=filename,
                date=datetime.now(),
                title=filename,
                content=content,
                content_summary="",
                main_topics=[],
                past_events=[],
                future_actions=[],
                participants=[],
                user_id=user_id,
                meeting_id=meeting_id,
                project_id=project_id,
                file_size=file_size
            )
            
            # Process with enhanced chunking
            chunks = self.chunk_document(meeting_doc)
            
            # Store document and chunks together using the existing add_document method
            self.vector_db.add_document(meeting_doc, chunks)
            
            # Update final status
            if status_id:
                self.vector_db.update_file_processing_status(status_id, 'completed', 
                                                           f'Successfully processed {len(chunks)} chunks')
            
            logger.info(f"Successfully processed {filename} with {len(chunks)} chunks")
            return {
                'success': True,
                'document_id': document_id,
                'chunks_created': len(chunks),
                'hash_id': hash_id
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            if status_id:
                self.vector_db.update_file_processing_status(status_id, 'failed', str(e))
            return {'success': False, 'error': str(e)}

    def process_files_batch_async(self, files: List[Dict], user_id: str, 
                                 project_id: str = None, meeting_id: str = None, 
                                 max_workers: int = 3, job_id: str = None) -> Dict:
        """Process multiple files asynchronously with threading"""
        if not files:
            return {'success': False, 'error': 'No files provided'}
        
        # Create upload job if not provided
        if job_id is None:
            job_id = self.vector_db.create_upload_job(user_id, len(files), project_id, meeting_id)
        
        # Create file processing status entries
        file_tasks = []
        for file_info in files:
            file_hash = self.vector_db.calculate_file_hash(file_info['path'])
            file_size = os.path.getsize(file_info['path'])
            
            # Check for duplicates first
            duplicate_info = self.vector_db.is_file_duplicate(file_hash, file_info['filename'], user_id)
            if duplicate_info:
                logger.info(f"Skipping duplicate file: {file_info['filename']}")
                continue
            
            status_id = self.vector_db.create_file_processing_status(job_id, file_info['filename'], 
                                                         file_size, file_hash)
            file_tasks.append({
                'file_path': file_info['path'],
                'filename': file_info['filename'],
                'status_id': status_id,
                'file_hash': file_hash
            })
        
        if not file_tasks:
            self.vector_db.update_job_status(job_id, 'completed', 0, 0)
            return {
                'success': True,
                'job_id': job_id,
                'message': 'All files were duplicates, no processing needed'
            }
        
        # Update job with actual file count after deduplication
        self.vector_db.update_job_status(job_id, 'processing')
        
        # Process files concurrently
        results = []
        processed_count = 0
        failed_count = 0
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_task = {
                executor.submit(
                    self.process_file_async,
                    task['file_path'],
                    task['filename'],
                    user_id,
                    project_id,
                    meeting_id,
                    job_id,
                    task['status_id']
                ): task for task in file_tasks
            }
            
            # Process completed tasks
            for future in as_completed(future_to_task):
                task = future_to_task[future]
                try:
                    result = future.result()
                    results.append({
                        'filename': task['filename'],
                        'result': result
                    })
                    
                    if result['success']:
                        processed_count += 1
                    else:
                        failed_count += 1
                        
                    # Update job progress
                    self.vector_db.update_job_status(job_id, 'processing', processed_count, failed_count)
                    
                except Exception as e:
                    logger.error(f"Error processing {task['filename']}: {e}")
                    failed_count += 1
                    results.append({
                        'filename': task['filename'],
                        'result': {'success': False, 'error': str(e)}
                    })
                    self.vector_db.update_job_status(job_id, 'processing', processed_count, failed_count)
        
        # Update final job status
        final_status = 'completed' if failed_count == 0 else 'partial' if processed_count > 0 else 'failed'
        self.vector_db.update_job_status(job_id, final_status, processed_count, failed_count)
        
        # Save vector index if any files were processed successfully
        if processed_count > 0:
            self.vector_db.save_index()
            logger.info(f"Saved FAISS index after processing {processed_count} documents")
        
        logger.info(f"Batch processing completed. Job: {job_id}, "
                   f"Processed: {processed_count}, Failed: {failed_count}")
        
        return {
            'success': True,
            'job_id': job_id,
            'processed_files': processed_count,
            'failed_files': failed_count,
            'total_files': len(file_tasks),
            'results': results
        }

def main():
    """Main function for Meeting Document AI System with OpenAI"""
    
    pass
    
    try:
        # Check OpenAI API key
        if not openai_api_key:
            pass
            return
        
        pass
        
        processor = EnhancedMeetingDocumentProcessor(chunk_size=1000, chunk_overlap=200)
        
        # Check if documents folder exists
        docs_folder = Path("meeting_documents")
        if not docs_folder.exists():
            docs_folder.mkdir(exist_ok=True)
            print("📁 Created 'meeting_documents' folder. Please add your meeting documents to this folder.")
            return
        
        # Check for documents
        doc_files = []
        for ext in ['.txt', '.docx', '.pdf']:
            doc_files.extend(docs_folder.glob(f"*{ext}"))
        
        if not doc_files:
            print("📁 No meeting documents found in the 'meeting_documents' folder.")
            print("📋 Supported formats: .txt, .docx, .pdf")
            print("📂 Please add your meeting documents to the 'meeting_documents' folder and run again.")
            return
        
        print(f"📄 Found {len(doc_files)} documents to process")
        
        # Check if vector database exists
        print("🔍 Checking existing vector database...")
        if processor.vector_db.index.ntotal == 0:
            print("🔄 Processing documents and building vector database...")
            processor.process_documents("meeting_documents")
        else:
            print(f"✅ Loaded existing vector database with {processor.vector_db.index.ntotal} vectors")
        
        # Show comprehensive statistics
        print("\n📊 System Statistics:")
        print("-" * 50)
        stats = processor.get_meeting_statistics()
        if "error" not in stats:
            print(f"📋 Total meetings: {stats.get('total_meetings', 0)}")
            print(f"🧩 Total chunks: {stats.get('total_chunks', 0)}")
            print(f"🔢 Vector index size: {stats.get('vector_index_size', 0)}")
            print(f"📏 Average chunk length: {stats.get('average_chunk_length', 0)} characters")
            
            if 'date_range' in stats:
                print(f"📅 Date range: {stats['date_range']['earliest']} to {stats['date_range']['latest']}")
            
            # AI Configuration details removed from display
        
        # Interactive query loop
        print("\n🎯 Interactive Query Session")
        print("-" * 50)
        print("💡 Now supports hundreds of documents with hybrid search!")
        print("🔍 Combines semantic similarity + keyword matching for better results")
        
        example_queries = [
            "What are the main topics from recent meetings?",
            "Tell me about the AI integration progress across all meetings",
            "What are our upcoming deadlines and action items?",
            "Summarize all migration plans discussed"
        ]
        
        for i, example in enumerate(example_queries[:3], 1):
            print(f"   {i}. {example}")
        print("   ... or ask anything about your meetings!")
        
        print(f"\n💬 Commands: 'quit'/'exit' to stop, 'stats' for statistics, 'help' for examples")
        
        while True:
            print("\n" + "-"*60)
            query = input("🤔 Your question: ").strip()
            
            if query.lower() in ['quit', 'exit', 'q']:
                print("👋 Thank you for using the Meeting Document AI System!")
                break
            elif query.lower() == 'stats':
                stats = processor.get_meeting_statistics()
                print("\n📊 Detailed Statistics:")
                print(json.dumps(stats, indent=2, default=str))
                continue
            elif query.lower() == 'help':
                print("\n📚 Example Queries (enhanced with hybrid search):")
                for i, example in enumerate(example_queries, 1):
                    print(f"   {i}. {example}")
                print("\n💭 More ideas:")
                print("   • Search across hundreds of documents instantly")
                print("   • Find specific keywords + semantic meaning")
                print("   • Get comprehensive answers from multiple meetings")
                print("   • Timeline analysis across large document sets")
                continue
            elif query.lower() == 'refresh':
                try:
                    processor.refresh_clients()
                    print("✅ OpenAI clients refreshed successfully")
                except Exception as e:
                    print(f"❌ Failed to refresh clients: {e}")
                continue
            elif not query:
                print("❓ Please enter a valid question.")
                continue
            
            print(f"\n🔍 Processing with hybrid search: '{query}'...")
            start_time = datetime.now()
            
            try:
                answer = processor.answer_query(query)
                processing_time = (datetime.now() - start_time).total_seconds()
                
                print(f"\n💬 Answer (processed in {processing_time:.2f}s):")
                print("-" * 60)
                print(answer)
                
            except Exception as query_error:
                print(f"❌ Error processing query: {query_error}")
                print("🔄 You can try rephrasing your question or check your OpenAI API key.")
    
    except Exception as e:
        logger.error(f"Critical error in main execution: {e}")
        print(f"❌ Critical Error: {e}")
        print("🔧 Please check your OpenAI API key and configuration.")
        print("💡 Make sure you have set OPENAI_API_KEY environment variable")

if __name__ == "__main__":
    main()